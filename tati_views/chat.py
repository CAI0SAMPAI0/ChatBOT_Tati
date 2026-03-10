"""
tati_views/chat.py — Teacher Tati · Chat principal + geração de arquivos.
"""

import json
import base64
from pathlib import Path

import streamlit as st
import streamlit.components.v1 as components

from database import (
    new_conversation, list_conversations, load_conversation,
    append_message, delete_conversation, cached_load_conversation,
)
from transcriber import transcribe_bytes
from file_reader import extract_file
from ui_helpers import (
    PROF_NAME, API_KEY,
    t, avatar_html, user_avatar_html, get_tati_mini_b64,
    render_audio_player, send_to_claude, do_logout,
)


def get_or_create_conv(username: str) -> str:
    if not st.session_state.conv_id:
        st.session_state.conv_id = new_conversation(username)
    return st.session_state.conv_id


# ══════════════════════════════════════════════════════════════════════════════
# GERAÇÃO DE ARQUIVOS (PDF / DOCX)
# ══════════════════════════════════════════════════════════════════════════════

def _intercept_file_generation(reply_text: str, username: str, conv_id: str) -> str:
    import re
    try:
        match = re.search(
            r'<<<GENERATE_FILE>>>\s*(\{.*?\})\s*<<<END_FILE>>>',
            reply_text, re.DOTALL,
        )
        if not match:
            append_message(username, conv_id, "assistant", reply_text)
            return reply_text

        meta     = json.loads(match.group(1))
        fmt      = meta.get("format","pdf").lower()
        title    = meta.get("title","Activity")
        content  = meta.get("content","")
        filename = meta.get("filename", f"activity.{fmt}")
        if not filename.endswith(f".{fmt}"):
            filename = f"{filename}.{fmt}"

        out_dir  = Path("data/generated")
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / filename

        if fmt == "pdf":
            _generate_pdf(title, content, out_path)
        else:
            _generate_docx(title, content, out_path)

        with open(out_path, "rb") as f:
            file_bytes = f.read()
        st.session_state["_pending_download"] = {
            "b64":      base64.b64encode(file_bytes).decode(),
            "filename": filename,
            "mime":     "application/pdf" if fmt == "pdf" else
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

        display_msg = (
            f"📎 Arquivo gerado: **{filename}**\n\n_{title}_\n\n"
            "Clique em **⬇ Baixar arquivo** abaixo para salvar."
        )
        append_message(username, conv_id, "assistant", display_msg, is_file=True)
        cached_load_conversation.clear()
        return display_msg

    except Exception as e:
        err = f"Desculpe, não consegui gerar o arquivo: {e}"
        append_message(username, conv_id, "assistant", err)
        cached_load_conversation.clear()
        return err


def _generate_pdf(title: str, content: str, out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles   import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units    import cm
    from reportlab.lib          import colors
    from reportlab.platypus     import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums    import TA_CENTER

    doc    = SimpleDocTemplate(str(out_path), pagesize=A4,
                               leftMargin=2.5*cm, rightMargin=2.5*cm,
                               topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    story  = []
    t_style = ParagraphStyle("t", parent=styles["Title"],  fontSize=18, spaceAfter=6,
                              textColor=colors.HexColor("#1a1a2e"), alignment=TA_CENTER)
    s_style = ParagraphStyle("s", parent=styles["Normal"], fontSize=9,  spaceAfter=14,
                              textColor=colors.HexColor("#888888"),  alignment=TA_CENTER)
    b_style = ParagraphStyle("b", parent=styles["Normal"], fontSize=11, leading=18, spaceAfter=8)

    story.append(Paragraph(title, t_style))
    story.append(Paragraph(f"Teacher {PROF_NAME}", s_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#f0a500")))
    story.append(Spacer(1, 0.4*cm))
    for line in content.split("\\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), b_style))
        else:
            story.append(Spacer(1, 0.2*cm))
    doc.build(story)


def _generate_docx(title: str, content: str, out_path: Path) -> None:
    from docx           import Document
    from docx.shared    import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.5)

    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    sub = doc.add_paragraph(f"Teacher {PROF_NAME}")
    sub.alignment           = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size      = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    for line in content.split("\\n"):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.style.font.size = Pt(11)
        else:
            doc.add_paragraph()
    doc.save(str(out_path))


# ══════════════════════════════════════════════════════════════════════════════
# PROCESSAMENTO DE ARQUIVOS ANEXADOS
# ══════════════════════════════════════════════════════════════════════════════

def _process_and_send_file(username, user, conv_id, raw, filename, extra_text=""):
    result = extract_file(raw, filename)
    kind, label = result["kind"], result["label"]

    if kind == "audio":
        with st.spinner("🔄 Transcrevendo áudio..."):
            text = transcribe_bytes(raw, suffix=Path(filename).suffix.lower(), language="en")
        if text.startswith("❌") or text.startswith("⚠️"):
            st.error(text); return False
        user_display = f"{extra_text}\n\n[Áudio transcrito: {text}]" if extra_text else text
        claude_msg   = f"{extra_text}\n\n[Áudio: '{filename}']\n{text}" if extra_text else \
                       f"[Áudio: '{filename}']\n{text}"
        append_message(username, conv_id, "user", user_display, audio=True)
        st.session_state.speaking = True
        try:   send_to_claude(username, user, conv_id, claude_msg)
        except Exception as e: st.error(f"❌ {e}")
        st.session_state.speaking = False
        return True

    elif kind == "text":
        extracted = result["text"]
        if extracted.startswith("❌"):  st.error(extracted);  return False
        if not extracted:               st.warning(f"Sem texto em '{filename}'."); return False
        preview      = extracted[:200].replace('\n', ' ')
        user_display = f"📄 [{label}: '{filename}'] — {preview}{'…' if len(extracted)>200 else ''}"
        if extra_text: user_display = f"{extra_text}\n\n{user_display}"
        claude_msg   = (f"📄 [{label}: '{filename}']\n\n{extracted}\n\n"
                        "Please help me understand this content — explain vocabulary, grammar, and key ideas.")
        if extra_text: claude_msg = f"{extra_text}\n\n{claude_msg}"
        append_message(username, conv_id, "user", user_display)
        st.session_state.speaking = True
        try:   send_to_claude(username, user, conv_id, claude_msg)
        except Exception as e: st.error(f"❌ {e}")
        st.session_state.speaking = False
        return True

    elif kind == "image":
        user_display = f"📸 [Imagem: '{filename}']"
        if extra_text: user_display = f"{extra_text}\n\n{user_display}"
        claude_msg = f"📸 [Imagem: '{filename}']\nPlease look at this image and help me learn English from it."
        if extra_text: claude_msg = f"{extra_text}\n\n{claude_msg}"
        append_message(username, conv_id, "user", user_display)
        st.session_state.speaking = True
        try:
            send_to_claude(username, user, conv_id, claude_msg,
                           image_b64=result["b64"], image_media_type=result["media_type"])
        except Exception as e: st.error(f"❌ {e}")
        st.session_state.speaking = False
        return True

    else:
        st.warning(f"⚠️ Formato '{label}' não suportado.")
        return False


# ══════════════════════════════════════════════════════════════════════════════
# CHAT PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

def show_chat() -> None:
    from tati_views.voice import show_voice_mode

    user     = st.session_state.user
    username = user["username"]
    profile  = user.get("profile", {})
    ui_lang  = profile.get("language", "pt-BR")
    conv_id  = get_or_create_conv(username)
    messages = cached_load_conversation(username, conv_id)
    speaking = st.session_state.speaking

    if st.session_state.voice_mode:
        show_voice_mode()
        return

    # Injeta cor de destaque
    _ac = profile.get("accent_color",      "#f0a500")
    _ub = profile.get("user_bubble_color", "#2d6a4f")
    _ab = profile.get("ai_bubble_color",   "#1a1f2e")
    components.html(f"""<!DOCTYPE html><html><head>
<style>html,body{{margin:0;padding:0;overflow:hidden;}}</style>
</head><body><script>
(function(){{
  function hexToRgb(h){{
    h=h.replace('#','');if(h.length===3)h=h[0]+h[0]+h[1]+h[1]+h[2]+h[2];
    var n=parseInt(h,16);return [(n>>16)&255,(n>>8)&255,n&255].join(',');
  }}
  function luminance(h){{
    h=h.replace('#','');if(h.length===3)h=h[0]+h[0]+h[1]+h[1]+h[2]+h[2];
    var n=parseInt(h,16);var r=(n>>16)&255,g=(n>>8)&255,b=n&255;
    return 0.299*r+0.587*g+0.114*b;
  }}
  var ac="{_ac}",ub="{_ub}",ab="{_ab}";
  var rgb=hexToRgb(ac);
  var r=window.parent.document.documentElement;
  r.style.setProperty('--accent-full',ac);
  r.style.setProperty('--accent-70','rgba('+rgb+',.7)');
  r.style.setProperty('--accent-40','rgba('+rgb+',.4)');
  r.style.setProperty('--accent-30','rgba('+rgb+',.3)');
  r.style.setProperty('--accent-15','rgba('+rgb+',.15)');
  r.style.setProperty('--bubble-bg','rgba('+rgb+',.12)');
  r.style.setProperty('--bubble-border','rgba('+rgb+',.3)');
  r.style.setProperty('--bubble-text','#e6edf3');
  r.style.setProperty('--user-bubble-bg',ub);
  r.style.setProperty('--user-bubble-text',luminance(ub)>128?'#111':'#e6edf3');
  r.style.setProperty('--ai-bubble-bg',ab);
  r.style.setProperty('--ai-bubble-text',luminance(ab)>128?'#111':'#e6edf3');
  r.style.setProperty('--ai-bubble-border','rgba('+hexToRgb(ab)+',.6)');
  // Para todo áudio ao trocar de conversa
  var par=window.parent;
  if(par){{
    par.document.querySelectorAll('audio').forEach(function(a){{a.pause();a.currentTime=0;}});
    par.document.querySelectorAll('iframe').forEach(function(f){{
      try{{
        f.contentDocument.querySelectorAll('audio').forEach(function(a){{a.pause();a.currentTime=0;}});
        f.contentDocument.querySelectorAll('#b').forEach(function(b){{b.textContent='\u25b6 Ouvir';}});
        if(f.contentWindow.speechSynthesis)f.contentWindow.speechSynthesis.cancel();
      }}catch(e){{}}
    }});
    if(par.speechSynthesis)par.speechSynthesis.cancel();
  }}
}})();
</script></body></html>""", height=1)

    # ── Sidebar ───────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown("""<style>
        section[data-testid="stSidebar"] { overflow:hidden; }
        section[data-testid="stSidebar"] > div:first-child {
            height:100vh;display:flex;flex-direction:column;
            padding:0!important;gap:0;
        }
        div.sidebar-footer { margin-top:auto; }
        </style>""", unsafe_allow_html=True)

        st.markdown(f"""<div style="padding:14px 14px 10px;border-bottom:1px solid #21262d;flex-shrink:0;">
            <div style="display:flex;align-items:center;gap:10px;">
                {avatar_html(40)}<div>
                <div style="font-weight:600;font-size:.88rem;">{PROF_NAME}</div>
                <div style="font-size:.68rem;color:#8b949e;"><span class="status-dot"></span>Online</div>
                </div></div></div>""", unsafe_allow_html=True)

        if st.button(t("new_conv", ui_lang), use_container_width=True, key="btn_new"):
            st.session_state.conv_id = new_conversation(username); st.rerun()
        if st.button(t("voice_mode", ui_lang), use_container_width=True, key="btn_voice"):
            st.session_state.voice_mode = True; st.rerun()

        st.markdown('<div style="font-size:.68rem;color:#8b949e;text-transform:uppercase;'
                    'letter-spacing:1px;padding:10px 4px 4px;">Conversas</div>',
                    unsafe_allow_html=True)

        convs = list_conversations(username)
        if not convs:
            st.markdown('<div style="font-size:.78rem;color:#8b949e;padding:6px 4px;">Nenhuma conversa ainda.</div>',
                        unsafe_allow_html=True)
        for c in convs:
            is_active = c["id"] == conv_id
            label     = ("▶ " if is_active else "") + c["title"]
            col_conv, col_del = st.columns([5, 1])
            with col_conv:
                if st.button(label, key=f"conv_{c['id']}", use_container_width=True,
                             help=f"📅 {c['date']} · 💬 {c['count']} msgs"):
                    st.session_state.conv_id = c["id"]; st.rerun()
            with col_del:
                if st.button("🗑", key=f"del_{c['id']}", help=t("delete_conv", ui_lang)):
                    delete_conversation(username, c["id"])
                    if st.session_state.conv_id == c["id"]:
                        st.session_state.conv_id = None
                    st.rerun()
            st.markdown(f'<div style="font-size:.62rem;color:#6e7681;margin:-10px 0 2px 6px;">'
                        f'📅 {c["date"]} · 💬 {c["count"]} msg</div>', unsafe_allow_html=True)

        # Rodapé sidebar
        user_msgs   = len([m for m in messages if m["role"] == "user"])
        uav_sidebar = user_avatar_html(username, size=34, fallback_emoji="🎓")
        st.markdown('<div class="sidebar-footer">', unsafe_allow_html=True)
        st.markdown("<hr style='border-color:#21262d;margin:8px 0 0'>", unsafe_allow_html=True)
        st.markdown(f"""<div style="padding:8px 12px;display:flex;align-items:center;gap:10px;">
            {uav_sidebar}
            <div style="flex:1;min-width:0;">
              <div style="font-weight:600;font-size:.82rem;white-space:nowrap;
                overflow:hidden;text-overflow:ellipsis;">{user['name'].split()[0]}</div>
              <div style="color:#8b949e;font-size:.68rem;">{user['level']} · {user_msgs} msgs</div>
            </div></div>""", unsafe_allow_html=True)

        if user["role"] == "professor":
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button(t("dashboard", ui_lang), use_container_width=True, key="btn_dash"):
                    st.session_state.page = "dashboard"; st.rerun()
            with col_b:
                if st.button(t("profile", ui_lang), use_container_width=True, key="btn_profile"):
                    st.session_state.page = "profile"; st.rerun()
            if st.button(t("logout", ui_lang), use_container_width=True, key="btn_sair"):
                do_logout(); st.rerun()
        else:
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button(t("profile", ui_lang), use_container_width=True, key="btn_profile"):
                    st.session_state.page = "profile"; st.rerun()
            with col_b:
                if st.button(t("logout", ui_lang), use_container_width=True, key="btn_sair"):
                    do_logout(); st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    # ── CSS do chat ───────────────────────────────────────────────────────────
    st.markdown("""<style>
[data-testid="stChatInput"] textarea {
    max-height:120px!important; min-height:44px!important; font-size:.88rem!important;
}
[data-testid="stChatInputContainer"] { padding:6px 10px!important; }
.main .block-container { padding-bottom:80px!important; }
.msg-row { display:flex; align-items:flex-end; gap:10px; margin:6px 0; }
.msg-row.user-row { flex-direction:row-reverse; justify-content:flex-start; }
.msg-row.bot-row  { flex-direction:row; }
.msg-row.user-row > div { display:flex; flex-direction:column; align-items:flex-end; }
.msg-row.bot-row > div  { display:flex; flex-direction:column; align-items:flex-start; }
.msg-bubble {
    padding:10px 15px; border-radius:18px;
    font-size:.88rem; line-height:1.6;
    word-break:normal; overflow-wrap:break-word; white-space:pre-wrap;
}
.msg-bubble.user {
    max-width:75%;
    background:var(--user-bubble-bg,#2d6a4f);
    color:var(--user-bubble-text,#d8f3dc);
    border-bottom-right-radius:4px;
}
.msg-bubble.bot {
    max-width:75%;
    background:var(--ai-bubble-bg,#1a1f2e);
    color:var(--ai-bubble-text,#e6edf3);
    border:1px solid var(--ai-bubble-border,#252d3d);
    border-bottom-left-radius:4px;
}
.msg-bubble.audio-msg { font-style:italic; opacity:.85; }
.msg-av { width:30px; height:30px; border-radius:50%; overflow:hidden; flex-shrink:0; margin-bottom:2px; }
.msg-av img { width:100%; height:100%; object-fit:cover; object-position:top; }
.msg-av .av-emoji { width:100%; height:100%; background:#1e2a3a; display:flex; align-items:center; justify-content:center; font-size:14px; }
.msg-time { font-size:.6rem; color:#4a5a6a; margin:2px 4px 0; text-align:right; }
.bot-row .msg-time { text-align:left; }
.msg-ouvir-row { padding:2px 0 0 40px; }
.msg-ouvir-btn {
    background:none; border:1px solid #30363d; border-radius:16px;
    color:#8b949e; font-size:.68rem; padding:2px 10px; cursor:pointer;
    transition:all .15s; white-space:nowrap; font-family:inherit;
}
.msg-ouvir-btn:hover,.msg-ouvir-btn.speaking { border-color:#f0a500; color:#f0a500; }
@media (max-width:768px) { .msg-bubble { max-width:88%!important; font-size:.82rem!important; } }
@media (max-width:480px) { .msg-bubble { max-width:94%!important; } }
</style>""", unsafe_allow_html=True)

    # ── Header ────────────────────────────────────────────────────────────────
    st.markdown(f"""<div class="prof-header">
        {avatar_html(56, speaking)}
        <div class="prof-info">
            <h1>{PROF_NAME}</h1>
            <p><span class="status-dot"></span>Online · {user['level']} · {user['focus']}</p>
        </div></div>""", unsafe_allow_html=True)

    # ── Histórico de mensagens ────────────────────────────────────────────────
    _tati_mini   = get_tati_mini_b64()
    tati_av_html = (
        f'<div class="msg-av" style="background:url({_tati_mini}) center top/cover no-repeat;"></div>'
        if _tati_mini else
        '<div class="msg-av"><div class="av-emoji">🧑‍🏫</div></div>'
    )

    st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
    for i, msg in enumerate(messages):
        content  = msg["content"].replace("\n","<br>")
        msg_time = msg.get("time","")

        if msg["role"] == "assistant":
            tts_b64 = msg.get("tts_b64","")
            is_file = msg.get("is_file",False)
            st.markdown(
                f'<div class="msg-row bot-row">{tati_av_html}'
                f'<div><div class="msg-bubble bot">{content}</div>'
                f'<div class="msg-time">{msg_time}</div></div></div>',
                unsafe_allow_html=True,
            )
            if tts_b64:
                components.html(render_audio_player(tts_b64, msg_time, f"msg_{i}_{conv_id}"),
                                height=44, scrolling=False)
            elif not is_file:
                clean_text = (msg["content"]
                    .replace("\\","").replace("`","")
                    .replace('"',"&quot;").replace("'","&#39;")
                    .replace("\n"," ").replace("\r","")
                    .replace("*","").replace("#",""))[:600]
                st.markdown(
                    f'<div class="msg-ouvir-row">'
                    f'<button class="msg-ouvir-btn" data-pav-tts data-text="{clean_text}">▶ Ouvir</button>'
                    f'</div>',
                    unsafe_allow_html=True,
                )
        else:
            is_audio = msg.get("audio",False)
            extra    = " audio-msg" if is_audio else ""
            st.markdown(
                f'<div class="msg-row user-row">'
                f'<div><div class="msg-bubble user{extra}">{content}</div>'
                f'<div class="msg-time">{msg_time}</div></div></div>',
                unsafe_allow_html=True,
            )
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Indicador "digitando" ─────────────────────────────────────────────────
    if st.session_state.get("speaking"):
        components.html("""<!DOCTYPE html><html><head>
<style>
*{margin:0;padding:0;box-sizing:border-box;}
html,body{background:transparent;overflow:hidden;font-family:'Sora',sans-serif;}
.typing-row{display:flex;align-items:center;gap:10px;padding:6px 0 4px 0;}
.av{width:30px;height:30px;border-radius:50%;background:#1e2a3a;
    display:flex;align-items:center;justify-content:center;font-size:14px;flex-shrink:0;}
.typing-bubble{display:flex;align-items:center;gap:8px;background:#1a1f2e;
    border:1px solid #252d3d;border-radius:18px;border-bottom-left-radius:4px;padding:8px 14px;}
.spin{color:#e05c2a;font-size:16px;line-height:1;animation:spinme 1.2s linear infinite;display:inline-block;}
@keyframes spinme{from{transform:rotate(0deg);}to{transform:rotate(360deg);}}
.typing-text{font-size:.75rem;color:#8b949e;font-style:italic;letter-spacing:.2px;}
</style></head><body>
<div class="typing-row">
  <div class="av">🧑‍🏫</div>
  <div class="typing-bubble">
    <span class="spin">✳</span>
    <span class="typing-text" id="msg">Pensando…</span>
  </div>
</div>
<script>
(function(){
  var msgs=[[0,"Pensando…"],[3,"Elaborando resposta…"],[7,"Demorando mais que o normal (tentativa 1)"],
            [14,"Demorando mais que o normal (tentativa 2)"],[22,"Demorando mais que o normal (tentativa 3)"]];
  var el=document.getElementById('msg'),start=Date.now();
  function update(){
    var sec=(Date.now()-start)/1000,txt=msgs[0][1];
    for(var i=0;i<msgs.length;i++){if(sec>=msgs[i][0])txt=msgs[i][1];}
    el.textContent=txt;setTimeout(update,800);
  }
  update();
})();
</script></body></html>""", height=52, scrolling=False)

    # ── Arquivo staged ────────────────────────────────────────────────────────
    staged = st.session_state.get("staged_file")
    if staged:
        staged_list = staged if isinstance(staged, list) else [staged]
        icons = {"audio":"🎵","text":"📄","image":"📸"}
        items_html = "".join(
            f'<span style="background:rgba(255,255,255,.06);border-radius:6px;'
            f'padding:3px 8px;font-size:.8rem;color:#e6edf3;">'
            f'{icons.get(f["kind"],"📎")} {f["name"]}</span>'
            for f in staged_list
        )
        st.markdown(f"""
<div style="background:rgba(240,165,0,.08);border:1px solid rgba(240,165,0,.25);
     border-radius:10px;padding:10px 14px;margin:6px 0;
     display:flex;align-items:center;justify-content:space-between;gap:10px;flex-wrap:wrap;">
  <div style="display:flex;gap:6px;flex-wrap:wrap;align-items:center;">
    {items_html}
    <span style="color:#8b949e;font-size:.75rem;">· {len(staged_list)} arquivo(s) anexado(s)</span>
  </div>
  <span style="font-size:.7rem;color:#f0a500;">↩ Digite uma mensagem ou envie</span>
</div>""", unsafe_allow_html=True)
        if st.button(t("remove_attachment", ui_lang), key="remove_staged"):
            st.session_state.staged_file      = None
            st.session_state.staged_file_name = None
            st.session_state.pop("_last_files_key", None)
            st.rerun()

    # ── Botão de download de arquivo gerado ───────────────────────────────────
    pending_dl = st.session_state.get("_pending_download")
    if pending_dl:
        b64_data = pending_dl["b64"]
        fname    = pending_dl["filename"]
        mime     = pending_dl["mime"]
        st.markdown(f"""
<div style="background:rgba(240,165,0,.08);border:1px solid rgba(240,165,0,.35);
     border-radius:10px;padding:10px 16px;margin:8px 0;display:flex;
     align-items:center;justify-content:space-between;gap:12px;">
  <span style="font-size:.85rem;color:#e6edf3;">📎 <b>{fname}</b> pronto para download</span>
  <a href="data:{mime};base64,{b64_data}" download="{fname}"
     style="background:linear-gradient(135deg,#f0a500,#e05c2a);color:#060a10;
     font-weight:700;font-size:.78rem;padding:6px 16px;border-radius:20px;
     text-decoration:none;white-space:nowrap;">⬇ Baixar arquivo</a>
</div>""", unsafe_allow_html=True)

    # ── Chat input ────────────────────────────────────────────────────────────
    prompt = st.chat_input(t("type_message", ui_lang))
    if prompt:
        if not API_KEY:
            st.error("Configure ANTHROPIC_API_KEY no .env"); st.stop()
        staged = st.session_state.get("staged_file")
        if staged:
            staged_list = staged if isinstance(staged, list) else [staged]
            for i, sf in enumerate(staged_list):
                _process_and_send_file(username, user, conv_id,
                                       sf["raw"], sf["name"], extra_text=prompt if i==0 else "")
            st.session_state.staged_file      = None
            st.session_state.staged_file_name = None
            st.session_state.pop("_last_files_key", None)
        else:
            append_message(username, conv_id, "user", prompt)
            st.session_state.speaking = True
            try:   send_to_claude(username, user, conv_id, prompt)
            except Exception as e: st.error(f"❌ {e}")
            st.session_state.speaking = False
        st.rerun()

    # ── Gravador de áudio nativo ──────────────────────────────────────────────
    audio_val = st.audio_input(
        " ", key=f"voice_input_{st.session_state.audio_key}",
        label_visibility="collapsed",
    )
    if audio_val and audio_val != st.session_state.get("_last_audio"):
        st.session_state["_last_audio"] = audio_val
        with st.spinner("Transcrevendo..."):
            txt = transcribe_bytes(audio_val.read(), ".wav", None)
        if txt and not txt.startswith("❌") and not txt.startswith("⚠️"):
            if not API_KEY: st.error("Configure ANTHROPIC_API_KEY"); st.stop()
            append_message(username, conv_id, "user", txt, audio=True)
            st.session_state.speaking = True
            try:   send_to_claude(username, user, conv_id, txt)
            except Exception as e: st.error(f"❌ {e}")
            st.session_state.speaking = False
            st.session_state.audio_key += 1
            st.rerun()
        elif txt:
            st.error(txt)

    # ── File uploader ─────────────────────────────────────────────────────────
    uploaded_list = st.file_uploader(
        "📎", key="file_upload", label_visibility="collapsed",
        accept_multiple_files=True,
        type=["mp3","wav","ogg","m4a","webm","flac",
              "pdf","doc","docx","txt","png","jpg","jpeg","webp"],
    )
    if uploaded_list:
        names_key = ",".join(sorted(f.name for f in uploaded_list))
        if names_key != st.session_state.get("_last_files_key"):
            st.session_state["_last_files_key"] = names_key
            staged_list = []
            for uf in uploaded_list:
                raw    = uf.read()
                result = extract_file(raw, uf.name)
                staged_list.append({"raw":raw,"name":uf.name,"kind":result["kind"],"result":result})
            st.session_state.staged_file      = staged_list
            st.session_state.staged_file_name = ", ".join(f["name"] for f in staged_list)
            st.rerun()

    # ── Botões "Ouvir" via SpeechSynthesis ───────────────────────────────────
    components.html("""<!DOCTYPE html><html><body><script>
(function(){
  var par=window.parent?window.parent.document:document;
  var cur=null;
  function initBtns(){
    par.querySelectorAll('[data-pav-tts]').forEach(function(btn){
      if(btn._pavInit)return;btn._pavInit=true;
      btn.addEventListener('click',function(){
        if(cur&&cur!==btn){speechSynthesis.cancel();cur.textContent='▶ Ouvir';cur.classList.remove('speaking');cur=null;}
        if(btn.classList.contains('speaking')){speechSynthesis.cancel();btn.textContent='▶ Ouvir';btn.classList.remove('speaking');cur=null;return;}
        var txt=btn.getAttribute('data-text')||'';
        var u=new SpeechSynthesisUtterance(txt);
        u.lang='en-US';u.rate=0.95;u.pitch=1.05;
        speechSynthesis.getVoices();
        setTimeout(function(){
          var vv=speechSynthesis.getVoices();
          var pick=vv.find(function(v){return v.lang==='en-US';})||vv.find(function(v){return v.lang.startsWith('en');});
          if(pick)u.voice=pick;
          u.onstart=function(){btn.textContent='⏹ Parar';btn.classList.add('speaking');cur=btn;};
          u.onend=u.onerror=function(){btn.textContent='▶ Ouvir';btn.classList.remove('speaking');cur=null;};
          speechSynthesis.cancel();speechSynthesis.speak(u);
        },80);
      });
    });
  }
  initBtns();
  new MutationObserver(initBtns).observe(par.body,{childList:true,subtree:true});
})();
</script></body></html>""", height=1)

    # ── Botões mic e clipe (static/pav_buttons.html) ──────────────────────────
    _btn_html = Path("static/pav_buttons.html")
    _btn_css  = Path("static/pav_buttons.css")
    if _btn_css.exists():
        st.markdown(f"<style>{_btn_css.read_text()}</style>", unsafe_allow_html=True)
    if _btn_html.exists():
        components.html(_btn_html.read_text(), height=1)
    else:
        st.warning("static/pav_buttons.html não encontrado")
