import os
import json
import base64
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()

import streamlit as st
import streamlit.components.v1 as components
import anthropic
from datetime import datetime

# ── Imports do banco e serviços ───────────────────────────────────────────────
from database import (
    init_db, authenticate, register_student, load_students,
    new_conversation, list_conversations, load_conversation,
    append_message, get_all_students_stats, delete_conversation,
    update_profile, update_password,
    create_session, validate_session, delete_session,
    save_user_avatar_db, get_user_avatar_db, remove_user_avatar_db, get_client, AVATAR_BUCKET   # sessões persistentes
)
from transcriber import transcribe_bytes
from tts import text_to_speech, tts_available
from file_reader import extract_file


# ── Font Awesome (ícones de anexo, etc.) ─────────────────────────────────────
st.markdown(
    '<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.1/css/all.min.css">',
    unsafe_allow_html=True
)

# ── JS: ajusta layout ao abrir/fechar sidebar automaticamente ─────────────────
components.html("""<!DOCTYPE html><html><head>
<style>html,body{margin:0;padding:0;overflow:hidden;}</style>
</head><body><script>
(function(){
    function inject(){
        var doc = window.parent.document;
        // Injeta <style> diretamente no head do pai
        if(!doc.getElementById('pav-toggle-style')){
            var s = doc.createElement('style');
            s.id = 'pav-toggle-style';
            s.textContent = '[data-testid="collapsedControl"]{position:fixed!important;top:10px!important;left:10px!important;z-index:99999!important;}';
            doc.head.appendChild(s);
        }
    }
    // Tenta imediatamente e depois com delay
    inject();
    setTimeout(inject, 500);
    setTimeout(inject, 1500);
    setInterval(inject, 3000);
})();
</script></body></html>""", height=1)

# login lendo o cookie

def _try_autologin_from_cookie() -> bool:
    """
    Tenta fazer login automático lendo o cookie pav_session do header HTTP.
    Retorna True se o login foi restaurado com sucesso.

    Requer Streamlit >= 1.31 (st.context.headers).
    Funciona em todos os browsers, inclusive Safari e iOS.
    """
    if st.session_state.logged_in:
        return True  # já logado

    # Tenta ler os headers HTTP da requisição atual
    try:
        headers = st.context.headers
        cookie_header = headers.get("Cookie", "") or headers.get("cookie", "")
    except AttributeError:
        return False

    if not cookie_header:
        return False

    # Extrai pav_session do header Cookie
    token = None
    for part in cookie_header.split(";"):
        part = part.strip()
        if part.startswith("pav_session="):
            import urllib.parse
            token = urllib.parse.unquote(part.split("=", 1)[1].strip())
            break

    if not token or len(token) < 10:
        return False

    # Valida o token no banco
    udata = validate_session(token)
    if not udata:
        return False

    # Resolve o username
    uname = udata.get("_resolved_username")
    if not uname:
        students = load_students()
        uname = next(
            (k for k, v in students.items() if v.get("password") == udata.get("password")),
            None
        )
    if not uname:
        return False

    # Restaura a sessão
    st.session_state.logged_in         = True
    st.session_state.user              = {"username": uname, **udata}
    st.session_state.page              = "dashboard" if udata.get("role") == "professor" else "chat"
    st.session_state.conv_id           = None
    st.session_state["_session_token"] = token
    return True

# ── Inicialização do banco de dados (SQLite) ──────────────────────────────────
init_db()

# ── Variáveis de ambiente ─────────────────────────────────────────────────────
API_KEY    = os.getenv("ANTHROPIC_API_KEY", "")
PHOTO_PATH = os.getenv("PROFESSOR_PHOTO", "assets/tati.png")
PROF_NAME  = os.getenv("PROFESSOR_NAME",  "Professor Avatar")

# ── Contador de áudio — força re-render do widget nativo ─────────────────────
if "audio_key" not in st.session_state:
    st.session_state.audio_key = 0



# ══════════════════════════════════════════════════════════════════════════════
# INTERNACIONALIZAÇÃO (i18n)
# ══════════════════════════════════════════════════════════════════════════════

_STRINGS = {
    "pt-BR": {
        "type_message":       "Digite uma mensagem…",
        "new_conv":           "➕  Nova conversa",
        "voice_mode":         "🎙️  Modo Conversa",
        "dashboard":          "📊 Painel",
        "profile":            "⚙️ Perfil",
        "logout":             "🚪 Sair",
        "delete_conv":        "Excluir conversa",
        "username":           "Usuário",
        "password":           "Senha",
        "full_name":          "Nome completo",
        "email":              "E-mail",
        "enter":              "Entrar",
        "create_account":     "Criar Conta",
        "save_general":       "💾 Salvar Alterações",
        "save_custom":        "💾 Salvar Personalização",
        "save_data":          "💾 Salvar Dados",
        "change_password":    "🔒 Alterar Senha",
        "remove_photo":       "🗑️ Remover foto",
        "remove_attachment":  "✕ Remover anexo",
        "close_voice":        "✕ Fechar Modo Voz",
        "close":              "✕ Fechar",
        "back":               "← Voltar ao Chat",
        "use_as_student":     "Usar como Aluno",
        "enter_chat":         "Entrar no Chat",
        "my_profile":         "⚙️ Meu Perfil",
        "interface_lang":     "Idioma da interface",
        "theme":              "Tema",
        "transcription_lang": "Idioma da transcrição (Whisper)",
        "tts_accent":         "Sotaque (TTS fallback)",
        "nickname":           "Apelido",
        "occupation":         "Ocupação",
        "english_level":      "Nível de inglês",
        "focus":              "Foco",
        "conv_tone":          "Tom das conversas",
        "ai_role":            "Papel da IA",
        "new_password":       "Nova senha",
        "confirm_password":   "Confirmar senha",
        "speaking":           "🗣 Falando...",
        "speaking_ai":        "🗣 Falando...",
        "listening":          "🎙 Ouvindo...",
        "processing":         "⏳ Processando...",
        "tap_to_speak":       "Toque no microfone para falar",
        "tap_to_record":      "Toque para gravar",
        "tap_to_stop":        "Toque para parar",
        "wait":               "Aguarde...",
        "conv_accent":        "Cor de destaque",
        "user_bubble":        "Balão do usuário",
        "ai_bubble":          "Balão da IA",
        "custom_instructions":"Instruções personalizadas para a IA",
        "custom_placeholder": "ex: Sempre me corrija quando eu errar o Past Simple.",
        "your_account":       "Sua Conta",
        "profile_photo":      "Foto de Perfil",
        "change_photo":       "Alterar foto — JPG, PNG ou WEBP (máx 15 MB)",
        "account_info":       "Informações da Conta",
        "username_label":     "Username:",
        "created_at":         "Conta criada em:",
        "change_pw_title":    "Alterar Senha",
        "pw_too_short":       "Senha muito curta.",
        "pw_mismatch":        "As senhas não coincidem.",
        "pw_changed":         "✅ Senha alterada!",
        "saved_ok":           "✅ Salvo!",
        "photo_saved":        "✅ Foto salva!",
        "photo_removed":      "Foto removida.",
        "photo_too_large":    "❌ Foto muito grande. Máximo 15 MB.",
        "appearance":         "Aparência",
        "voice_settings":     "Voz",
        "about_you":          "Sobre Você",
        "ai_style_title":     "Estilo da IA",
        "conversations":      "Conversas",
        "no_convs":           "Nenhuma conversa ainda.",
        "msgs":               "msgs",
        "online":             "● Online",
        "students":           "Alunos",
        "messages":           "Mensagens",
        "corrections":        "Correções",
        "active_today":       "Ativos Hoje",
        "no_students":        "Nenhum aluno ainda.",
        "level":              "Nível",
        "last_access":        "Último Acesso",
        "fill_fields":        "Preencha todos os campos.",
        "invalid_email":      "E-mail inválido.",
        "pw_min":             "Senha muito curta (mínimo 6 caracteres).",
        "wrong_credentials":  "Usuário ou senha incorretos.",
        "back_dashboard":     "← Voltar ao Dashboard",
        "back_chat":          "← Voltar ao Chat",
    },
    "en-US": {
        "type_message":       "Type a message…",
        "new_conv":           "➕  New conversation",
        "voice_mode":         "🎙️  Voice Mode",
        "dashboard":          "📊 Dashboard",
        "profile":            "⚙️ Profile",
        "logout":             "🚪 Logout",
        "delete_conv":        "Delete conversation",
        "username":           "Username",
        "password":           "Password",
        "full_name":          "Full name",
        "email":              "E-mail",
        "enter":              "Sign In",
        "create_account":     "Create Account",
        "save_general":       "💾 Save Changes",
        "save_custom":        "💾 Save Customization",
        "save_data":          "💾 Save Data",
        "change_password":    "🔒 Change Password",
        "remove_photo":       "🗑️ Remove photo",
        "remove_attachment":  "✕ Remove attachment",
        "close_voice":        "✕ Close Voice Mode",
        "close":              "✕ Close",
        "back":               "← Back to Chat",
        "use_as_student":     "Use as Student",
        "enter_chat":         "Enter Chat",
        "my_profile":         "⚙️ My Profile",
        "interface_lang":     "Interface language",
        "theme":              "Theme",
        "transcription_lang": "Transcription language (Whisper)",
        "tts_accent":         "Accent (TTS fallback)",
        "nickname":           "Nickname",
        "occupation":         "Occupation",
        "english_level":      "English level",
        "focus":              "Focus",
        "conv_tone":          "Conversation tone",
        "ai_role":            "AI role",
        "new_password":       "New password",
        "confirm_password":   "Confirm password",
        "speaking":           "🗣 Speaking...",
        "speaking_ai":        "🗣 Speaking...",
        "listening":          "🎙 Listening...",
        "processing":         "⏳ Processing...",
        "tap_to_speak":       "Tap the mic to speak",
        "tap_to_record":      "Tap to record",
        "tap_to_stop":        "Tap to stop",
        "wait":               "Please wait...",
        "conv_accent":        "Accent color",
        "user_bubble":        "User bubble",
        "ai_bubble":          "AI bubble",
        "custom_instructions":"Custom AI instructions",
        "custom_placeholder": "e.g. Always correct me when I get the Past Simple wrong.",
        "your_account":       "Your Account",
        "profile_photo":      "Profile Photo",
        "change_photo":       "Change photo — JPG, PNG or WEBP (max 15 MB)",
        "account_info":       "Account Information",
        "username_label":     "Username:",
        "created_at":         "Account created on:",
        "change_pw_title":    "Change Password",
        "pw_too_short":       "Password too short.",
        "pw_mismatch":        "Passwords do not match.",
        "pw_changed":         "✅ Password changed!",
        "saved_ok":           "✅ Saved!",
        "photo_saved":        "✅ Photo saved!",
        "photo_removed":      "Photo removed.",
        "photo_too_large":    "❌ Photo too large. Maximum 15 MB.",
        "appearance":         "Appearance",
        "voice_settings":     "Voice",
        "about_you":          "About You",
        "ai_style_title":     "AI Style",
        "conversations":      "Conversations",
        "no_convs":           "No conversations yet.",
        "msgs":               "msgs",
        "online":             "● Online",
        "students":           "Students",
        "messages":           "Messages",
        "corrections":        "Corrections",
        "active_today":       "Active Today",
        "no_students":        "No students yet.",
        "level":              "Level",
        "last_access":        "Last Access",
        "fill_fields":        "Please fill in all fields.",
        "invalid_email":      "Invalid email address.",
        "pw_min":             "Password too short (minimum 6 characters).",
        "wrong_credentials":  "Incorrect username or password.",
        "back_dashboard":     "← Back to Dashboard",
        "back_chat":          "← Back to Chat",
    },
    "en-UK": {
        "type_message":       "Type a message…",
        "new_conv":           "➕  New conversation",
        "voice_mode":         "🎙️  Voice Mode",
        "dashboard":          "📊 Dashboard",
        "profile":            "⚙️ Profile",
        "logout":             "🚪 Log out",
        "delete_conv":        "Delete conversation",
        "username":           "Username",
        "password":           "Password",
        "full_name":          "Full name",
        "email":              "E-mail",
        "enter":              "Sign In",
        "create_account":     "Create Account",
        "save_general":       "💾 Save Changes",
        "save_custom":        "💾 Save Customisation",
        "save_data":          "💾 Save Data",
        "change_password":    "🔒 Change Password",
        "remove_photo":       "🗑️ Remove photo",
        "remove_attachment":  "✕ Remove attachment",
        "close_voice":        "✕ Close Voice Mode",
        "close":              "✕ Close",
        "back":               "← Back to Chat",
        "use_as_student":     "Use as Student",
        "enter_chat":         "Enter Chat",
        "my_profile":         "⚙️ My Profile",
        "interface_lang":     "Interface language",
        "theme":              "Theme",
        "transcription_lang": "Transcription language (Whisper)",
        "tts_accent":         "Accent (TTS fallback)",
        "nickname":           "Nickname",
        "occupation":         "Occupation",
        "english_level":      "English level",
        "focus":              "Focus",
        "conv_tone":          "Conversation tone",
        "ai_role":            "AI role",
        "new_password":       "New password",
        "confirm_password":   "Confirm password",
        "speaking":           "🗣 Speaking...",
        "speaking_ai":        "🗣 Speaking...",
        "listening":          "🎙 Listening...",
        "processing":         "⏳ Processing...",
        "tap_to_speak":       "Tap the mic to speak",
        "tap_to_record":      "Tap to record",
        "tap_to_stop":        "Tap to stop",
        "wait":               "Please wait...",
        "conv_accent":        "Accent colour",
        "user_bubble":        "User bubble",
        "ai_bubble":          "AI bubble",
        "custom_instructions":"Custom AI instructions",
        "custom_placeholder": "e.g. Always correct me when I get the Past Simple wrong.",
        "your_account":       "Your Account",
        "profile_photo":      "Profile Photo",
        "change_photo":       "Change photo — JPG, PNG or WEBP (max 15 MB)",
        "account_info":       "Account Information",
        "username_label":     "Username:",
        "created_at":         "Account created on:",
        "change_pw_title":    "Change Password",
        "pw_too_short":       "Password too short.",
        "pw_mismatch":        "Passwords do not match.",
        "pw_changed":         "✅ Password changed!",
        "saved_ok":           "✅ Saved!",
        "photo_saved":        "✅ Photo saved!",
        "photo_removed":      "Photo removed.",
        "photo_too_large":    "❌ Photo too large. Maximum 15 MB.",
        "appearance":         "Appearance",
        "voice_settings":     "Voice",
        "about_you":          "About You",
        "ai_style_title":     "AI Style",
        "conversations":      "Conversations",
        "no_convs":           "No conversations yet.",
        "msgs":               "msgs",
        "online":             "● Online",
        "students":           "Students",
        "messages":           "Messages",
        "corrections":        "Corrections",
        "active_today":       "Active Today",
        "no_students":        "No students yet.",
        "level":              "Level",
        "last_access":        "Last Access",
        "fill_fields":        "Please fill in all fields.",
        "invalid_email":      "Invalid e-mail address.",
        "pw_min":             "Password too short (minimum 6 characters).",
        "wrong_credentials":  "Incorrect username or password.",
        "back_dashboard":     "← Back to Dashboard",
        "back_chat":          "← Back to Chat",
    },
}

def t(key: str, lang: str = "pt-BR") -> str:
    """Retorna string traduzida. Fallback: pt-BR."""
    return _STRINGS.get(lang, _STRINGS["pt-BR"]).get(key, _STRINGS["pt-BR"].get(key, key))


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE IMAGEM / AVATAR
# ══════════════════════════════════════════════════════════════════════════════

def get_photo_b64() -> str | None:
    """Lê a foto da professora e devolve como data-URI base64."""
    p = Path(PHOTO_PATH)
    if p.exists():
        ext  = p.suffix.lower().replace(".", "")
        mime = "jpeg" if ext in ("jpg", "jpeg") else ext
        return f"data:image/{mime};base64,{base64.b64encode(p.read_bytes()).decode()}"
    return None

PHOTO_B64 = get_photo_b64()

# ── Cache da foto mini da Tati (evita re-leitura de disco a cada render) ──────
@st.cache_data(show_spinner=False)
def get_tati_mini_b64() -> str:
    """Lê a foto da Tati uma única vez e reutiliza em todo o app."""
    for _p in [Path("assets/tati.png"), Path("assets/tati.jpg"),
               Path(__file__).parent / "assets" / "tati.png",
               Path(__file__).parent / "assets" / "tati.jpg"]:
        if _p.exists():
            _ext  = _p.suffix.lstrip(".").lower()
            _mime = "jpeg" if _ext in ("jpg", "jpeg") else _ext
            return f"data:image/{_mime};base64,{base64.b64encode(_p.read_bytes()).decode()}"
    return get_photo_b64() or ""

# ── Cache dos 4 frames do avatar animado do modo voz ─────────────────────────
@st.cache_data(show_spinner=False)
def get_avatar_frames() -> dict:
    """Carrega os frames do avatar animado uma única vez."""
    _base = Path(__file__).parent
    def _load(candidates):
        for p in candidates:
            p = Path(p)
            if p.exists():
                return f"data:image/png;base64,{base64.b64encode(p.read_bytes()).decode()}"
        return ""
    return {
        # ── Fala da IA ──────────────────────────────────────────────────────
        "normal":     _load([_base/"assets"/"avatar_tati_normal.png",     "assets/avatar_tati_normal.png"]),
        "meio":       _load([_base/"assets"/"avatar_tati_meio.png",       "assets/avatar_tati_meio.png"]),
        "aberta":     _load([_base/"assets"/"avatar_tati_aberta.png",     "assets/avatar_tati_aberta.png"]),
        "bem_aberta": _load([_base/"assets"/"avatar_tati_bem_aberta.png", "assets/avatar_tati_bem_aberta.png"]),
        # ── Estados especiais ────────────────────────────────────────────────
        "ouvindo":    _load([_base/"assets"/"avatar_tati_ouvindo.png",    "assets/avatar_tati_ouvindo.png"]),
        "piscando":   _load([_base/"assets"/"tati_piscando.png",          "assets/tati_piscando.png"]),
        "surpresa":   _load([_base/"assets"/"tati_surpresa.png",          "assets/tati_surpresa.png"]),
    }

# ── Avatares individuais dos alunos ───────────────────────────────────────────
def get_user_avatar_b64(username: str, _bust: int = 0) -> str | None:
    """Busca foto do usuário direto do banco, sem cache."""
    result = get_user_avatar_db(username)
    if not result:
        return None
    raw, mime = result
    return f"data:{mime};base64,{base64.b64encode(raw).decode()}"

# alias usado internamente
_get_avatar = lambda u: get_user_avatar_b64(u, _bust=st.session_state.get("_avatar_v", 0))

def _avatar_circle_html(b64: str | None, size: int, border: str = "#8800f0") -> str:
    """Retorna HTML de avatar circular — foto, sem_foto.png ou ícone FA."""
    if not b64:
        for _p in [Path("assets/sem_foto.png"), Path(__file__).parent / "assets" / "sem_foto.png"]:
            if _p.exists():
                b64 = f"data:image/png;base64,{base64.b64encode(_p.read_bytes()).decode()}"
                break
    if b64:
        return (
            f'<div style="width:{size}px;height:{size}px;border-radius:50%;'
            f'background:url({b64}) center/cover no-repeat;'
            f'border:2px solid {border};flex-shrink:0;"></div>'
        )
    icon_px = int(size * 0.50)
    return (
        f'<div style="width:{size}px;height:{size}px;border-radius:50%;'
        f'background:linear-gradient(135deg,#1e2a3a,#2a3a50);'
        f'display:flex;align-items:center;justify-content:center;'
        f'border:2px solid #1e2a3a;flex-shrink:0;">'
        f'<i class="fa-duotone fa-solid fa-user-graduate" '
        f'style="font-size:{icon_px}px;--fa-primary-color:#f0a500;--fa-secondary-color:#c87800;--fa-secondary-opacity:0.6;"></i>'
        f'</div>'
    )

def save_user_avatar(username: str, raw: bytes, suffix: str) -> None:
    """Salva a foto de perfil no Supabase Storage."""
    suffix = suffix.lower().lstrip(".")
    mime   = "image/jpeg" if suffix in ("jpg", "jpeg") else f"image/{suffix}"
    save_user_avatar_db(username, raw, mime)
    _bump_avatar_version()

def remove_user_avatar(username: str) -> None:
    """Remove a foto de perfil do Supabase Storage."""
    remove_user_avatar_db(username)
    _bump_avatar_version()

def _bump_avatar_version() -> None:
    """Incrementa o contador de versão do avatar para forçar re-fetch."""
    st.session_state["_avatar_v"] = st.session_state.get("_avatar_v", 0) + 1

def user_avatar_html(username: str, size: int = 36, **_) -> str:
    """Retorna HTML de avatar circular do usuário."""
    return _avatar_circle_html(get_user_avatar_b64(username, _bust=st.session_state.get("_avatar_v", 0)), size)

def avatar_html(size: int = 52, speaking: bool = False) -> str:
    """Avatar da professora com anel de 'speaking' animado."""
    cls   = "speaking" if speaking else ""
    photo = PHOTO_B64
    if photo:
        # background-image no div: sem <img>, sem JS, sem flash
        return (
            f'<div class="avatar-wrap {cls}" style="'
            f'width:{size}px;height:{size}px;border-radius:50%;flex-shrink:0;'
            f'background:url({photo}) center top/cover no-repeat;'
            f'position:relative;overflow:hidden;">'
            f'<div class="avatar-ring"></div></div>'
        )
    return (
        f'<div class="avatar-circle {cls}" '
        f'style="width:{size}px;height:{size}px;font-size:{int(size*.48)}px">🧑‍🏫</div>'
    )


# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURAÇÃO STREAMLIT
# ══════════════════════════════════════════════════════════════════════════════

st.set_page_config(page_title=f"{PROF_NAME} · English", page_icon=str(Path(PHOTO_PATH)) if Path(PHOTO_PATH).exists() else "🎓", layout="wide")

def load_css(path: str) -> None:
    p = Path(path)
    if p.exists():
        css = p.read_text(encoding="utf-8")
        st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)

load_css("styles/style.css")

# ── CSS responsivo global ──────────────────────────────────────────────────────
st.markdown("""<style>
/* Remove overlay de escurecimento durante rerun */
[data-testid="stAppViewBlockContainer"] { opacity: 1 !important; }
div[data-stale="true"] { opacity: 1 !important; transition: none !important; }
div[data-stale="false"] { opacity: 1 !important; transition: none !important; }
/* Esconde spinner/loading global do Streamlit */
.stSpinner, [data-testid="stSpinner"],
div[class*="StatusWidget"], div[class*="stStatusWidget"] { display: none !important; }
/* Remove o fade/dim que o Streamlit aplica durante rerun */
.stApp > div { opacity: 1 !important; }
iframe[title="streamlit_loading"] { display: none !important; }
section[data-testid="stMain"] > div {
    transition: all .25s ease;
    max-width: 100% !important;
}
/* Evita flash de imagem desestilizada */
img { max-width: 100%; }
.avatar-wrap img, .avatar-ring img { display: block; will-change: opacity; }
/* Esconde conteúdo do st.markdown durante re-render */
[data-testid="stMarkdownContainer"] img {
    opacity: 1;
    transition: opacity .15s;
}
.main .block-container {
    max-width: 100% !important;
    padding-left: clamp(10px, 2vw, 40px) !important;
    padding-right: clamp(10px, 2vw, 40px) !important;
    padding-top: 1rem !important;
}
section[data-testid="stSidebar"] {
    min-width: 220px !important;
    max-width: 340px !important;
}
div[data-testid="stButton"] button {
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
}
/* Esconde file uploader nativo imediatamente — o JS depois move o clipe */
[data-testid="stFileUploader"] {
    position: fixed !important;
    bottom: -999px !important;
    left: -9999px !important;
    opacity: 0 !important;
    width: 1px !important;
    height: 1px !important;
    pointer-events: none !important;
    overflow: hidden !important;
}
[data-testid="stFileUploader"] input[type="file"] {
    pointer-events: auto !important;
}
/* Esconde audio input nativo imediatamente — o JS cria botão mic falso na chat bar */
[data-testid="stAudioInput"] {
    position: fixed !important;
    bottom: -9999px !important;
    left: -9999px !important;
    opacity: 0 !important;
    pointer-events: none !important;
    width: 1px !important;
    height: 1px !important;
    overflow: hidden !important;
}
[data-testid="stAudioInput"] button {
    pointer-events: auto !important;
}
@media (max-width: 1024px) {
    .main .block-container { padding-left: 12px !important; padding-right: 12px !important; }
}
@media (max-width: 768px) {
    section[data-testid="stSidebar"] { min-width: 0 !important; }
    .main .block-container {
        padding-left: 8px !important;
        padding-right: 8px !important;
        padding-bottom: 80px !important;
    }
    div.bubble { max-width: 92% !important; font-size: .82rem !important; }
    div.prof-header h1 { font-size: 1rem !important; }
    div.prof-header { padding: 10px 8px !important; }
    .bav-s, .bav-u { display: none !important; }
}
</style>""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PROMPT DO SISTEMA
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = f"""You are a digital avatar of an English teacher called {PROF_NAME} — warm, witty, very intelligent and encouraging. You help adults speak English with more confidence, over 25 years of experience, Advanced English Hunter College NY, and passionate about teaching.
Students: teenagers (Beginner/Pre-Intermediate) and adults focused on Business/News.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BILINGUAL POLICY (VERY IMPORTANT)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
The student's messages may arrive in English, Portuguese, or a mix.
Adapt your language policy according to the student's level:

BEGINNER / PRE-INTERMEDIATE:
  • Student writes/speaks in Portuguese → Fully acceptable. Respond in simple English
    AND provide the Portuguese translation of key words in parentheses.
    Example: "Great question! The word is 'homework' (tarefa)."
  • Student mixes PT and EN → Celebrate the English parts, gently supply the missing
    English for the Portuguese parts. Never make them feel bad for using Portuguese.
  • Always end your reply with an easy, encouraging question in English.
  • Provide Portuguese support freely when they seem lost or frustrated.

INTERMEDIATE:
  • Respond primarily in English. Use Portuguese ONLY to clarify a specific word
    or resolve a genuine comprehension block — keep it brief.
  • If the student writes in Portuguese, acknowledge briefly in English and invite
    them to try saying the same thing in English:
    "I understood! Now, how would you say that in English?"
  • Encourage them to push further; celebrate every English sentence they produce.

ADVANCED / BUSINESS ENGLISH:
  • Respond exclusively in English.
  • If the student writes in Portuguese, reply in English and say something like:
    "Let's keep it in English — you've got this!"
  • You may add a brief Portuguese gloss ONLY for highly technical or idiomatic
    terms where the meaning is genuinely ambiguous.

TRANSLATION REQUESTS (any level):
  • When the student asks "como se diz X?", "what does Y mean?", or similar,
    always provide the translation + an example sentence in English.
  • For Beginners/Pre-Intermediate: also include the Portuguese example.

TEACHING STYLE:
- Neuro-learning: guide students to discover errors. Never just give the answer.
  Example: "he go" → "What ending do we add for he/she/it?"
- Sandwich: 1) Validate 2) Guide with question 3) Encourage.
- SHORT conversational responses. Bold grammar points when appropriate.
- End responses with ONE engaging question.
- NEVER use emojis in your responses. No exceptions. Plain text only, always.

RULES:
- Simple English. Teens→Fortnite/Netflix/TikTok/Movies and series refs. Adults→LinkedIn/news/geopolitics.
- Portuguese → briefly acknowledge, when asked to speak Portuguese, speak, but switch to English.
- NEVER start a conversation uninvited. Wait for the student to speak first.
- NEVER use emojis. Not a single one. Ever.

ACTIVITY GENERATION:
- When asked to create exercises, worksheets or activities, generate complete, well-structured content.
- Support: fill-in-the-blank, multiple choice, reading comprehension, dialogue writing,
  grammar drills, vocabulary lists, translation, error correction, etc.
- When the student asks for a FILE (PDF, Word/DOCX), respond ONLY with a special JSON block
  on its own line, no other text, in this exact format:
  <<<GENERATE_FILE>>>
  {{"format":"pdf","filename":"activity.pdf","title":"Exercise Title","content":"Full content here with \\n for line breaks"}}
  <<<END_FILE>>>
  Use "pdf" or "docx" as format. Put ALL the exercise content inside "content".
  The system will intercept this and generate the real file for download."""


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE — valores padrão na primeira execução
# ══════════════════════════════════════════════════════════════════════════════

_defaults = {
    "logged_in":        False,
    "user":             None,
    "page":             "chat",
    "speaking":         False,
    "conv_id":          None,
    "voice_mode":       False,
    "staged_file":      None,   # arquivo anexado aguardando envio
    "staged_file_name": None,
}
for k, v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

_session_token = st.session_state.get("_session_token", "")
if _session_token and st.session_state.logged_in:
    # Mantém o token na URL para que o reload funcione
    if st.query_params.get("s") != _session_token:
        st.query_params["s"] = _session_token

if not st.session_state.logged_in:
    _s = st.query_params.get("s", "")
    if _s and len(_s) > 10:
        _udata = validate_session(_s)
        if _udata:
            _uname = _udata.get("_resolved_username") or next(
                (k for k, v in load_students().items() if v["password"] == _udata["password"]),
                None
            )
            if _uname:
                st.session_state.logged_in         = True
                st.session_state.user              = {"username": _uname, **_udata}
                st.session_state.page              = "dashboard" if _udata["role"] == "professor" else "chat"
                st.session_state.conv_id           = None
                st.session_state["_session_token"] = _s
        else:
            # Token inválido — limpa
            st.query_params.pop("s", None)

# ══════════════════════════════════════════════════════════════════════════════
# SESSÃO PERSISTENTE — funções de salvar/limpar no localStorage
# ══════════════════════════════════════════════════════════════════════════════

def js_save_session(token: str) -> None:
    """Salva token no localStorage E cookie (30 dias) para persistência total."""
    components.html(
        f"""<!DOCTYPE html><html><head>
<style>html,body{{margin:0;padding:0;overflow:hidden;}}</style>
</head><body><script>
        (function() {{
            var t = '{token}';
            try {{ window.parent.localStorage.setItem('pav_session', t); }} catch(e) {{}}
            try {{ localStorage.setItem('pav_session', t); }} catch(e) {{}}
            try {{
                var exp = new Date(Date.now()+2592000000).toUTCString();
                window.parent.document.cookie = 'pav_session='+encodeURIComponent(t)+';expires='+exp+';path=/;SameSite=Lax';
            }} catch(e) {{}}
            try {{
                var exp2 = new Date(Date.now()+2592000000).toUTCString();
                document.cookie = 'pav_session='+encodeURIComponent(t)+';expires='+exp2+';path=/;SameSite=Lax';
            }} catch(e) {{}}
        }})();
        </script></body></html>""",
        height=1,
    )

def js_clear_session() -> None:
    components.html(
        """<!DOCTYPE html><html><head>
<style>html,body{margin:0;padding:0;overflow:hidden;}</style>
</head><body><script>
        (function() {
            try { window.parent.localStorage.removeItem('pav_session'); } catch(e) {}
            try { localStorage.removeItem('pav_session'); } catch(e) {}
            try { window.parent.document.cookie='pav_session=;expires=Thu,01 Jan 1970 00:00:00 GMT;path=/'; } catch(e) {}
            try { document.cookie='pav_session=;expires=Thu,01 Jan 1970 00:00:00 GMT;path=/'; } catch(e) {}
        })();
        </script></body></html>""",
        height=1,
    )

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE CONVERSA / CLAUDE
# ══════════════════════════════════════════════════════════════════════════════

def get_or_create_conv(username: str) -> str:
    """Retorna o conv_id ativo ou cria uma nova conversa."""
    if not st.session_state.conv_id:
        st.session_state.conv_id = new_conversation(username)
    return st.session_state.conv_id


@st.cache_data(ttl=10, show_spinner=False)
def cached_load_conversation(username: str, conv_id: str) -> list:
    """Cache de 10s do histórico — evita re-leitura do banco a cada render."""
    return load_conversation(username, conv_id)


def send_to_claude(username: str, user: dict, conv_id: str,
                   text: str, image_b64: str = None, image_media_type: str = None) -> str:
    """
    Envia a mensagem ao Claude Haiku, obtém a resposta,
    gera TTS opcional e salva no banco.
    Retorna o texto da resposta.
    """
    client  = anthropic.Anthropic(api_key=API_KEY)
    #context = f"\n\nStudent: Name={user['name']}, Level={user['level']}, Focus={user['focus']}."
    context = (
       f"\n\nStudent profile — Name: {user['name']} | "
       f"Level: {user['level']} | Focus: {user['focus']} | "
       f"Native language: Brazilian Portuguese.\n"
       f"Apply the bilingual policy for level '{user['level']}' as instructed."
   )


    # Monta histórico da conversa para a API
    # Carrega direto do banco (sem cache) para garantir que a mensagem recém salva apareça
    msgs     = load_conversation(username, conv_id)
    api_msgs = [
        {"role": "user" if m["role"] == "user" else "assistant", "content": m["content"]}
        for m in msgs
    ]

    # Se a última mensagem do histórico não for do usuário com o texto atual,
    # adiciona explicitamente para evitar erro por cache desatualizado
    if not api_msgs or api_msgs[-1]["role"] != "user" or api_msgs[-1]["content"] != text:
        api_msgs.append({"role": "user", "content": text})

    # Adiciona imagem à última mensagem do usuário, se houver
    if image_b64 and image_media_type and api_msgs and api_msgs[-1]["role"] == "user":
        api_msgs[-1]["content"] = [
            {"type": "image", "source": {"type": "base64", "media_type": image_media_type, "data": image_b64}},
            {"type": "text",  "text": text}
        ]

    # Mais tokens para pedidos de atividade/arquivo
    is_activity = any(w in text.lower() for w in [
        "pdf", "word", "docx", "atividade", "exercício", "exercicio",
        "worksheet", "activity", "exercise", "generate", "criar arquivo", "crie um", "make a", "gere um"
    ])
    max_tok = 2000 if is_activity else 400

    resp       = client.messages.create(
        model="claude-haiku-4-5", max_tokens=max_tok,
        system=SYSTEM_PROMPT + context, messages=api_msgs
    )
    reply_text = resp.content[0].text

    # Remove emojis da resposta (garantia extra além do system prompt)
    import re as _re
    reply_text = _re.sub(
        r'[\U00010000-\U0010ffff'
        r'\U0001F300-\U0001F9FF'
        r'\u2600-\u26FF\u2700-\u27BF'
        r'\U0001FA00-\U0001FA6F\U0001FA70-\U0001FAFF'
        r'\u200d\ufe0f]',
        '', reply_text
    ).strip()

    # Verifica se a IA quer gerar um arquivo (PDF/DOCX)
    if "<<<GENERATE_FILE>>>" in reply_text:
        return _intercept_file_generation(reply_text, username, conv_id)

    # Gera áudio TTS e armazena no banco junto com a mensagem
    tts_b64_str = None
    if tts_available():
        audio_bytes = text_to_speech(reply_text)
        if audio_bytes:
            tts_b64_str = base64.b64encode(audio_bytes).decode()
            st.session_state["_tts_audio"] = tts_b64_str

    append_message(username, conv_id, "assistant", reply_text, tts_b64=tts_b64_str)
    cached_load_conversation.clear()  # invalida cache para próximo render
    return reply_text


def _intercept_file_generation(reply_text: str, username: str, conv_id: str) -> str:
    """
    Intercepta o bloco <<<GENERATE_FILE>>> na resposta da IA
    e cria o arquivo PDF ou DOCX real para download.
    """
    import re

    try:
        match = re.search(
            r'<<<GENERATE_FILE>>>\s*(\{.*?\})\s*<<<END_FILE>>>',
            reply_text, re.DOTALL
        )
        if not match:
            append_message(username, conv_id, "assistant", reply_text)
            return reply_text

        meta     = json.loads(match.group(1))
        fmt      = meta.get("format", "pdf").lower()
        title    = meta.get("title", "Activity")
        content  = meta.get("content", "")
        filename = meta.get("filename", f"activity.{fmt}")
        if not filename.endswith(f".{fmt}"):
            filename = f"{filename}.{fmt}"

        out_dir  = Path("data/generated")
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / filename

        if fmt == "pdf":
            _generate_pdf(title, content, out_path)
        else:
            _generate_docx(title, content, out_path)

        # Armazena para exibir o botão de download no chat
        with open(out_path, "rb") as f:
            file_bytes = f.read()
        st.session_state["_pending_download"] = {
            "b64":      base64.b64encode(file_bytes).decode(),
            "filename": filename,
            "mime":     "application/pdf" if fmt == "pdf" else
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

        display_msg = (
            f"📎 Arquivo gerado: **{filename}**\n\n_{title}_\n\n"
            "Clique em **⬇ Baixar arquivo** abaixo para salvar."
        )
        append_message(username, conv_id, "assistant", display_msg, is_file=True)
        cached_load_conversation.clear()
        return display_msg

    except Exception as e:
        err = f"Desculpe, não consegui gerar o arquivo: {e}"
        append_message(username, conv_id, "assistant", err)
        cached_load_conversation.clear()
        return err


def _generate_pdf(title: str, content: str, out_path: Path) -> None:
    """Gera PDF profissional com ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles   import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units    import cm
    from reportlab.lib          import colors
    from reportlab.platypus     import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums    import TA_CENTER

    doc    = SimpleDocTemplate(str(out_path), pagesize=A4,
                               leftMargin=2.5*cm, rightMargin=2.5*cm,
                               topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    story  = []

    t_style = ParagraphStyle("t", parent=styles["Title"],   fontSize=18, spaceAfter=6,
                              textColor=colors.HexColor("#1a1a2e"), alignment=TA_CENTER)
    s_style = ParagraphStyle("s", parent=styles["Normal"],  fontSize=9,  spaceAfter=14,
                              textColor=colors.HexColor("#888888"),  alignment=TA_CENTER)
    b_style = ParagraphStyle("b", parent=styles["Normal"],  fontSize=11, leading=18, spaceAfter=8)

    story.append(Paragraph(title, t_style))
    story.append(Paragraph(f"Teacher {PROF_NAME}", s_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#f0a500")))
    story.append(Spacer(1, 0.4*cm))

    for line in content.split("\\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), b_style))
        else:
            story.append(Spacer(1, 0.2*cm))

    doc.build(story)


def _generate_docx(title: str, content: str, out_path: Path) -> None:
    """Gera DOCX profissional com python-docx."""
    from docx           import Document
    from docx.shared    import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5); sec.right_margin  = Cm(2.5)

    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    sub = doc.add_paragraph(f"Teacher {PROF_NAME}")
    sub.alignment       = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size      = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    for line in content.split("\\n"):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.style.font.size = Pt(11)
        else:
            doc.add_paragraph()

    doc.save(str(out_path))


# ══════════════════════════════════════════════════════════════════════════════
# PLAYER DE ÁUDIO TTS — mini player inline nos balões da IA
# ══════════════════════════════════════════════════════════════════════════════

def render_audio_player(tts_b64: str, msg_time: str, player_id: str) -> str:
    return f"""<!DOCTYPE html><html><head>
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
html,body{{background:transparent;font-family:'Sora',sans-serif;overflow:hidden;}}
.player{{display:flex;align-items:center;gap:8px;padding:3px 0;flex-wrap:wrap;}}
.tl{{font-size:.62rem;color:#8b949e;font-family:'JetBrains Mono',monospace;flex-shrink:0;}}
.pb{{background:none;border:1px solid #30363d;border-radius:20px;color:#f0a500;font-size:.75rem;
     padding:2px 10px;cursor:pointer;transition:background .15s,border-color .15s;white-space:nowrap;flex-shrink:0;}}
.pb:hover{{background:rgba(240,165,0,.12);border-color:#f0a500;}}
.pw{{flex:1;min-width:60px;height:3px;background:#30363d;border-radius:2px;cursor:pointer;}}
.pf{{height:100%;background:linear-gradient(90deg,#f0a500,#e05c2a);border-radius:2px;width:0%;transition:width .1s linear;pointer-events:none;}}
.sw{{display:flex;align-items:center;gap:3px;flex-shrink:0;}}
.sb{{background:none;border:1px solid #30363d;border-radius:4px;color:#8b949e;font-size:.65rem;
     padding:1px 5px;cursor:pointer;transition:all .15s;}}
.sb:hover,.sb.on{{border-color:#f0a500;color:#f0a500;background:rgba(240,165,0,.08);}}
.vw{{display:flex;align-items:center;gap:4px;flex-shrink:0;}}
.vi{{font-size:.75rem;cursor:pointer;color:#8b949e;}}
.vs{{width:50px;height:3px;accent-color:#f0a500;cursor:pointer;}}
</style></head><body>
<div class="player">
  <span class="tl">{msg_time}</span>
  <button class="pb" id="b">▶ Ouvir</button>
  <div class="pw" id="pw"><div class="pf" id="pf"></div></div>
  <div class="sw" id="sw">
    <span style="font-size:.6rem;color:#8b949e;">vel:</span>
    <button class="sb" data-r="0.75">0.75×</button>
    <button class="sb on" data-r="1">1×</button>
    <button class="sb" data-r="1.25">1.25×</button>
    <button class="sb" data-r="1.5">1.5×</button>
  </div>
  <div class="vw">
    <span class="vi" id="vi">🔊</span>
    <input type="range" class="vs" id="vs" min="0" max="1" step="0.05" value="1">
  </div>
</div>
<script>
(function(){{
  var audio=new Audio('data:audio/mpeg;base64,{tts_b64}');
  audio.preload='metadata';
  var b=document.getElementById('b'),pf=document.getElementById('pf'),
      pw=document.getElementById('pw'),vs=document.getElementById('vs'),
      vi=document.getElementById('vi'),sw=document.getElementById('sw');

  // Desbloqueia AudioContext no iOS/Android
  function unlockAudio(){{
    try{{
      var AC=window.AudioContext||window.webkitAudioContext;
      if(!AC)return;
      var ctx=new AC();
      var buf=ctx.createBuffer(1,1,22050);
      var src=ctx.createBufferSource();
      src.buffer=buf; src.connect(ctx.destination); src.start(0);
      if(ctx.state==='suspended')ctx.resume();
    }}catch(e){{}}
    // Também notifica o parent
    try{{
      if(window.parent&&window.parent.pavUnlockAudio)window.parent.pavUnlockAudio();
    }}catch(e){{}}
  }}

  b.onclick=function(){{
    unlockAudio();
    if(!audio.paused){{
      audio.pause(); b.textContent='▶ Ouvir';
    }}else{{
      var p=audio.play();
      if(p!==undefined){{
        p.then(function(){{b.textContent='⏸ Pausar';}})
         .catch(function(err){{
           console.warn('play() bloqueado:', err);
           b.textContent='▶ Ouvir';
           setTimeout(function(){{
             audio.play()
               .then(function(){{b.textContent='⏸ Pausar';}})
               .catch(function(){{}});
           }},300);
         }});
      }}
    }}
  }};

  audio.onended=function(){{b.textContent='▶ Ouvir';pf.style.width='0%';}};
  audio.ontimeupdate=function(){{
    if(audio.duration)pf.style.width=(audio.currentTime/audio.duration*100)+'%';
  }};
  pw.onclick=function(e){{
    var r=pw.getBoundingClientRect();
    if(audio.duration)audio.currentTime=((e.clientX-r.left)/r.width)*audio.duration;
  }};
  sw.querySelectorAll('.sb').forEach(function(btn){{
    btn.onclick=function(){{
      sw.querySelectorAll('.sb').forEach(function(x){{x.classList.remove('on');}});
      this.classList.add('on');
      audio.playbackRate=parseFloat(this.dataset.r);
    }};
  }});
  vs.oninput=function(){{
    audio.volume=parseFloat(vs.value);
    vi.textContent=audio.volume===0?'🔇':audio.volume<0.5?'🔉':'🔊';
  }};
  vi.onclick=function(){{
    if(audio.volume>0){{audio._v=audio.volume;audio.volume=0;vs.value=0;vi.textContent='🔇';}}
    else{{audio.volume=audio._v||1;vs.value=audio.volume;vi.textContent='🔊';}}
  }};
}})();
</script></body></html>"""

# ══════════════════════════════════════════════════════════════════════════════
# TELA DE LOGIN / REGISTRO
# ══════════════════════════════════════════════════════════════════════════════

def show_login() -> None:
    photo_src = get_photo_b64() or ""

    # Remove TODO o padding/margin do Streamlit e aplica roxo nos botões
    st.markdown("""<style>
[data-testid='stSidebar']{display:none!important;}
#MainMenu,footer,header{display:none!important;}
[data-testid="stToolbar"],[data-testid="stHeader"],[data-testid="stDecoration"]{display:none!important;}
.stApp{background:#060a10!important;}
section[data-testid="stMain"],
section[data-testid="stMain"]>div,
.main .block-container{
    padding:0!important;margin:0!important;
    max-width:100%!important;width:100%!important;
}
/* Botões de aba secundários */
div[data-testid="stButton"]>button{
    border-radius:10px!important;font-weight:600!important;
    border:1px solid #2a2a4a!important;
    background:transparent!important;color:#6b7280!important;
}
/* Botão de aba ativo (primary) */
div[data-testid="stButton"]>button[kind="primary"],
div[data-testid="stButton"]>button[data-testid="baseButton-primary"]{
    background:linear-gradient(135deg,#6c3fc5,#8b5cf6)!important;
    border-color:#7c4dcc!important;color:#fff!important;
    box-shadow:0 0 14px rgba(139,92,246,.35)!important;
}
/* Botão submit do form */
div[data-testid="stFormSubmitButton"]>button{
    background:linear-gradient(135deg,#6c3fc5,#8b5cf6)!important;
    border:1px solid #7c4dcc!important;color:#fff!important;
    border-radius:10px!important;font-weight:700!important;
    box-shadow:0 0 14px rgba(139,92,246,.3)!important;
}
div[data-testid="stFormSubmitButton"]>button:hover{
    background:linear-gradient(135deg,#7c4dcc,#9d6ff7)!important;
    box-shadow:0 0 22px rgba(139,92,246,.5)!important;
}
/* Iframes fantasma invisíveis */
iframe[height="1"]{
    position:fixed!important;opacity:0!important;
    pointer-events:none!important;bottom:0!important;left:0!important;
}
/* Centraliza tudo verticalmente */
section[data-testid="stMain"]>div>div>div{
    display:flex!important;
    flex-direction:column!important;
    align-items:center!important;
}
div[data-testid="stVerticalBlock"]{
    width:100%!important;
    max-width:420px!important;
    margin:0 auto!important;
    padding:0 16px!important;
}
</style>""", unsafe_allow_html=True)

    # ── Auto-login via localStorage ──────────────────────────────────────────
    components.html("""<!DOCTYPE html><html><head>
<style>html,body{margin:0;padding:0;overflow:hidden;}</style>
</head><body><script>
(function(){
    function readToken(){
        try{var s=window.parent.localStorage.getItem('pav_session');if(s&&s.length>10)return s;}catch(e){}
        try{var s2=localStorage.getItem('pav_session');if(s2&&s2.length>10)return s2;}catch(e){}
        try{var m=window.parent.document.cookie.split(';').map(function(c){return c.trim();})
            .find(function(c){return c.startsWith('pav_session=');});
            if(m){var v=decodeURIComponent(m.split('=')[1]);if(v&&v.length>10)return v;}}catch(e){}
        return '';
    }
    var val=readToken();
    if(!val)return;
    var url=new URL(window.parent.location.href);
    if(url.searchParams.get('s')!==val){
        url.searchParams.set('s',val);
        window.parent.location.replace(url.toString());
    }
})();
</script></body></html>""", height=1)

    if "_login_tab" not in st.session_state:
        st.session_state["_login_tab"] = "login"

    # ── Avatar HTML ──────────────────────────────────────────────────────────
    if photo_src:
        av_html = (
            f'<img class="av" src="{photo_src}" alt="{PROF_NAME}" '
            f'onerror="this.style.display=\'none\';'
            f'document.getElementById(\'avE\').style.display=\'flex\';">'
            f'<div class="av-emoji" id="avE" style="display:none">&#129489;&#8203;&#127979;</div>'
        )
    else:
        av_html = '<div class="av-emoji" id="avE">&#129489;&#8203;&#127979;</div>'

    # ── Card visual (avatar + nome) ──────────────────────────────────────────
    components.html(f"""<!DOCTYPE html>
<html><head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1,maximum-scale=1">
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@400;700;800&display=swap');
*{{box-sizing:border-box;margin:0;padding:0;}}
html,body{{
    background:#060a10;font-family:'Sora',sans-serif;
    width:100%;height:100%;overflow:hidden;
    display:flex;align-items:center;justify-content:center;
}}
.card{{
    background:linear-gradient(180deg,#0f1824,#0a1020);
    border:1px solid #1a2535;border-radius:24px;
    padding:28px 24px 20px;width:100%;
    box-shadow:0 24px 64px rgba(0,0,0,.7);
    display:flex;flex-direction:column;align-items:center;
}}
.av{{
    width:90px;height:90px;border-radius:50%;
    object-fit:cover;object-position:top center;
    border:2.5px solid #8b5cf6;
    box-shadow:0 0 0 6px rgba(139,92,246,.12),0 0 28px rgba(139,92,246,.25);
    display:block;margin-bottom:12px;
}}
.av-emoji{{
    width:90px;height:90px;border-radius:50%;
    background:linear-gradient(135deg,#6c3fc5,#8b5cf6);
    display:flex;align-items:center;justify-content:center;
    font-size:38px;margin-bottom:12px;
}}
h2{{
    font-size:1.35rem;font-weight:800;text-align:center;margin:0 0 3px;
    background:linear-gradient(135deg,#8b5cf6 30%,#c084fc 100%);
    -webkit-background-clip:text;-webkit-text-fill-color:transparent;
}}
p{{font-size:.7rem;color:#3a4e5e;text-align:center;}}
</style></head><body>
<div class="card">
    {av_html}
    <h2>{PROF_NAME}</h2>
    <p>Voice English Coach</p>
</div>
</body></html>""", height=220, scrolling=False)

    # ── Mensagens de feedback ────────────────────────────────────────────────
    login_err = st.session_state.pop("_login_err", "")
    reg_err   = st.session_state.pop("_reg_err",   "")
    reg_ok    = st.session_state.pop("_reg_ok",    False)
    reg_name  = st.session_state.pop("_reg_name",  "")
    if login_err: st.error(f"❌ {login_err}")
    if reg_err:   st.error(f"❌ {reg_err}")
    if reg_ok:    st.success(f"✅ Conta criada! Bem-vindo(a), {reg_name}!")

    # ── Abas ─────────────────────────────────────────────────────────────────
    tab = st.session_state["_login_tab"]
    c1, c2 = st.columns(2)
    with c1:
        if st.button(t("enter"), use_container_width=True, key="tab_login",
                     type="primary" if tab == "login" else "secondary"):
            st.session_state["_login_tab"] = "login"
            st.rerun()
    with c2:
        if st.button(t("create_account"), use_container_width=True, key="tab_reg",
                     type="primary" if tab == "reg" else "secondary"):
            st.session_state["_login_tab"] = "reg"
            st.rerun()

    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

    # ── Formulários ──────────────────────────────────────────────────────────
    if tab == "login":
        with st.form("form_login", clear_on_submit=True):
            u = st.text_input(t("username"), placeholder="seu.usuario")
            p = st.text_input(t("password"), type="password", placeholder="••••••••")
            if st.form_submit_button(t("enter"), use_container_width=True):
                if not u or not p:
                    st.session_state["_login_err"] = "Preencha todos os campos."
                    st.rerun()
                else:
                    user = authenticate(u, p)
                    if user:
                        real_u = user.get("_resolved_username", u.lower())
                        st.session_state.update(
                            logged_in=True,
                            user={"username": real_u, **user},
                            page="dashboard" if user["role"] == "professor" else "voice",
                            conv_id=None,
                        )
                        token = create_session(real_u)
                        st.session_state["_session_token"] = token
                        st.session_state["_session_saved"] = True
                        js_save_session(token)
                        st.rerun()
                    else:
                        st.session_state["_login_err"] = "Usuário ou senha incorretos."
                        st.rerun()
    else:
        with st.form("form_reg", clear_on_submit=True):
            rn  = st.text_input(t("full_name"),  placeholder="João Silva")
            re_ = st.text_input(t("email"),      placeholder="joao@email.com")
            ru  = st.text_input(t("username"),   placeholder="joao.silva")
            rp  = st.text_input("Senha", type="password", placeholder="mínimo 6 caracteres")
            if st.form_submit_button(t("create_account"), use_container_width=True):
                if not rn or not re_ or not ru or not rp:
                    st.session_state["_reg_err"] = "Preencha todos os campos."
                    st.rerun()
                elif "@" not in re_:
                    st.session_state["_reg_err"] = "E-mail inválido."
                    st.rerun()
                elif len(rp) < 6:
                    st.session_state["_reg_err"] = "Senha muito curta (mínimo 6)."
                    st.rerun()
                else:
                    ok, msg = register_student(ru, rn, rp, email=re_)
                    if ok:
                        st.session_state["_reg_ok"]    = True
                        st.session_state["_reg_name"]  = rn
                        st.session_state["_login_tab"] = "login"
                        st.rerun()
                    else:
                        st.session_state["_reg_err"] = msg
                        st.rerun()

    st.markdown(
        f'<p style="text-align:center;font-size:.6rem;color:#1a2535;margin-top:14px;">'
        f'2025 © {PROF_NAME}</p>',
        unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# PERFIL DO USUÁRIO
# ══════════════════════════════════════════════════════════════════════════════

def show_profile() -> None:
    """Página de configurações: aparência, personalização e dados da conta."""
    user     = st.session_state.user
    username = user["username"]
    profile  = user.get("profile", {})
    ui_lang  = profile.get("language", "pt-BR")

    # Garante que o file uploader de foto seja visível nesta página
    st.markdown("""<style>
[data-testid="stFileUploader"] {
    position: static !important; top: auto !important; left: auto !important;
    width: auto !important; height: auto !important;
    overflow: visible !important; opacity: 1 !important;
    pointer-events: auto !important;
}
[data-testid="stFileUploaderDropzone"] { display: flex !important; visibility: visible !important; }
</style>""", unsafe_allow_html=True)

    _ac = profile.get("accent_color", "#f0a500")
    _ub = profile.get("user_bubble_color", "#2d6a4f")
    _ab = profile.get("ai_bubble_color", "#1a1f2e")
    components.html(f"""<!DOCTYPE html><html><head>
<style>html,body{{margin:0;padding:0;overflow:hidden;}}</style>
</head><body><script>
(function(){{
  function hexToRgb(h){{
    h=h.replace('#','');
    if(h.length===3) h=h[0]+h[0]+h[1]+h[1]+h[2]+h[2];
    var n=parseInt(h,16);
    return [(n>>16)&255,(n>>8)&255,n&255].join(',');
  }}
  function luminance(h){{
    h=h.replace('#','');
    if(h.length===3) h=h[0]+h[0]+h[1]+h[1]+h[2]+h[2];
    var n=parseInt(h,16);
    var r=(n>>16)&255, g=(n>>8)&255, b=n&255;
    return 0.299*r + 0.587*g + 0.114*b;
  }}
  var ac="{_ac}", ub="{_ub}", ab="{_ab}";
  var rgb=hexToRgb(ac);
  var r=window.parent.document.documentElement;
  r.style.setProperty('--accent-full',ac);
  r.style.setProperty('--accent-70','rgba('+rgb+',.7)');
  r.style.setProperty('--accent-40','rgba('+rgb+',.4)');
  r.style.setProperty('--accent-30','rgba('+rgb+',.3)');
  r.style.setProperty('--accent-15','rgba('+rgb+',.15)');
  r.style.setProperty('--bubble-bg','rgba('+rgb+',.12)');
  r.style.setProperty('--bubble-border','rgba('+rgb+',.3)');
  r.style.setProperty('--bubble-text','#e6edf3');
  r.style.setProperty('--user-bubble-bg', ub);
  r.style.setProperty('--user-bubble-text', luminance(ub)>128 ? '#111' : '#e6edf3');
  r.style.setProperty('--ai-bubble-bg', ab);
  r.style.setProperty('--ai-bubble-text', luminance(ab)>128 ? '#111' : '#e6edf3');
  r.style.setProperty('--ai-bubble-border', 'rgba('+hexToRgb(ab)+', .6)');
}})();
</script></body></html>""", height=1)

    st.markdown("## ⚙️ Configurações do Perfil")
    st.markdown("---")

    is_prof    = user.get("role") == "professor"
    level_opts = ["Beginner", "Pre-Intermediate", "Intermediate", "Business English", "Advanced", "Native"]
    focus_opts = ["General Conversation","Business English","Travel","Academic",
                  "Pronunciation","Grammar","Vocabulary","Exam Prep"]

    def safe_index(lst, val, default=0):
        try:    return lst.index(val)
        except: return default

    tab_geral, tab_pers, tab_conta = st.tabs(["🎨 Geral", "🧠 Personalização", "👤 Conta"])

    # ── Aba Geral ─────────────────────────────────────────────────────────────
    with tab_geral:
        st.markdown("### Aparência")
        col1, col2 = st.columns(2)
        with col1:
            lang  = st.selectbox(t("interface_lang", ui_lang), ["pt-BR", "en-US", "en-UK"],
                index=safe_index(["pt-BR", "en-US", "en-UK"], profile.get("language", "pt-BR")), key="pf_lang")
        with col2:
            accent = st.color_picker("Cor de destaque (anel / botões)",
                value=profile.get("accent_color", "#f0a500"), key="pf_accent")

        col5, col6 = st.columns(2)
        with col5:
            user_bubble_color = st.color_picker("Balão do usuário",
                value=profile.get("user_bubble_color", "#2d6a4f"), key="pf_user_bubble")
        with col6:
            ai_bubble_color = st.color_picker("Balão da IA",
                value=profile.get("ai_bubble_color", "#1a1f2e"), key="pf_ai_bubble")

        st.markdown("### Voz")
        col3, col4 = st.columns(2)
        with col3:
            voice_lang = st.selectbox(
       t("transcription_lang", ui_lang),
       ["auto (pt+en)", "en", "pt", "es", "fr", "de"],
       index=safe_index(["auto (pt+en)", "en", "pt", "es", "fr", "de"],
                        profile.get("voice_lang", "auto (pt+en)")),
       key="pf_vlang"
   )
        with col4:
            speech_lang = st.selectbox(t("tts_accent", ui_lang), ["en-US", "en-UK", "pt-BR"],
                index=safe_index(["en-US", "en-UK", "pt-BR"], profile.get("speech_lang", "en-US")), key="pf_slang")

        if st.button(t("save_general", ui_lang), key="save_geral"):
            update_profile(username, {"language": lang,
                "accent_color": accent,
                "user_bubble_color": user_bubble_color,
                "ai_bubble_color": ai_bubble_color,
                "voice_lang": voice_lang, "speech_lang": speech_lang})
            u = load_students().get(username, {})
            # Atualiza session_state para refletir cor imediatamente
            st.session_state.user = {"username": username, **u}
            st.success("✅ Settings saved!")

    # ── Aba Personalização ────────────────────────────────────────────────────
    with tab_pers:
        st.markdown("### Sobre Você")
        col1, col2 = st.columns(2)
        with col1:
            nickname   = st.text_input(t("nickname", ui_lang), value=profile.get("nickname", ""), key="pf_nick")
            occupation = st.text_input(t("occupation", ui_lang), value=profile.get("occupation", ""),
                placeholder="ex: Professora, Desenvolvedor", key="pf_occ")
        with col2:
            level = st.selectbox(t("english_level", ui_lang), level_opts,
                index=safe_index(level_opts, user.get("level", "Beginner")), key="pf_level")
            focus = st.selectbox(t("focus", ui_lang), focus_opts,
                index=safe_index(focus_opts, user.get("focus", "General Conversation")), key="pf_focus")

        if not is_prof:
            st.markdown("### Estilo da IA")
            col3, col4 = st.columns(2)
            ai_style_opts = ["Warm & Encouraging", "Formal & Professional", "Fun & Casual", "Strict & Direct"]
            ai_tone_opts  = ["Teacher", "Conversation Partner", "Tutor", "Business Coach"]
            with col3:
                ai_style = st.selectbox(t("conv_tone", ui_lang), ai_style_opts,
                    index=safe_index(ai_style_opts, profile.get("ai_style", "Warm & Encouraging")), key="pf_aistyle")
            with col4:
                ai_tone = st.selectbox(t("ai_role", ui_lang), ai_tone_opts,
                    index=safe_index(ai_tone_opts, profile.get("ai_tone", "Teacher")), key="pf_aitone")
            custom = st.text_area("Instruções personalizadas para a IA",
                value=profile.get("custom_instructions", ""),
                placeholder="ex: Sempre me corrija quando eu errar o Past Simple.",
                height=100, key="pf_custom")
        else:
            ai_style = profile.get("ai_style", "Warm & Encouraging")
            ai_tone  = profile.get("ai_tone",  "Teacher")
            custom   = profile.get("custom_instructions", "")

        if st.button(t("save_custom", ui_lang), key="save_pers"):
            update_profile(username, {"nickname": nickname, "occupation": occupation,
                "ai_style": ai_style, "ai_tone": ai_tone, "custom_instructions": custom,
                "level": level, "focus": focus})
            u = load_students().get(username, {})
            st.session_state.user = {"username": username, **u}
            st.success("✅ Perfil salvo!")

    # ── Aba Conta ─────────────────────────────────────────────────────────────
    with tab_conta:
        st.markdown("### 📸 Foto de Perfil")

        msg = st.session_state.pop("_photo_msg", None)
        if msg == "saved":   st.success("✅ Foto salva!")
        elif msg == "removed": st.success("Foto removida.")

        cur_avatar = get_user_avatar_b64(username, _bust=st.session_state.get("_avatar_v", 0))
        MAX_BYTES  = 15 * 1024 * 1024

        col_av, col_btns = st.columns([1, 3])
        with col_av:
            st.markdown(
                _avatar_circle_html(cur_avatar, size=88) +
                '<div style="height:8px"></div>',
                unsafe_allow_html=True)
        with col_btns:
            photo_file = st.file_uploader(
                "Alterar foto — JPG, PNG ou WEBP (máx 15 MB)",
                type=["jpg", "jpeg", "png", "webp"], key="pf_photo_upload")
            if photo_file:
                file_id = f"{photo_file.name}::{photo_file.size}"
                if st.session_state.get("_last_photo_saved") != file_id:
                    raw_photo = photo_file.read()
                    if len(raw_photo) > MAX_BYTES:
                        st.error("❌ Foto muito grande. Máximo 15 MB.")
                    else:
                        suffix = Path(photo_file.name).suffix.lstrip(".")
                        save_user_avatar(username, raw_photo, suffix)
                        st.session_state["_last_photo_saved"] = file_id
                        _bump_avatar_version()
                        st.session_state["_photo_msg"] = "saved"
                        st.rerun()
                        if cur_avatar:
                            if st.button(t("remove_photo", ui_lang), key="pf_remove_photo"):
                                remove_user_avatar(username)
                                _bump_avatar_version()
                                #get_user_avatar_b64.clear()
                                st.session_state.user.get("profile", {}).pop("avatar_v", None)
                                st.session_state.user["profile"] = {
                                    k: v for k, v in st.session_state.user.get("profile", {}).items()
                                    if k != "avatar_v"
                                }
                                st.session_state.pop("_last_photo_saved", None)
                                st.session_state["_photo_msg"] = "removed"
                                st.rerun()

        st.markdown("---")
        st.markdown("### Informações da Conta")
        col1, col2 = st.columns(2)
        with col1:
            full_name = st.text_input(t("full_name", ui_lang), value=user.get("name", ""), key="pf_fname")
        with col2:
            email = st.text_input(t("email", ui_lang), value=user.get("email", ""), key="pf_email")

        st.markdown(f"**Username:** `{username}`")
        st.markdown(f"**Conta criada em:** {user.get('created_at', '')[:10]}")

        if st.button(t("save_data", ui_lang), key="save_conta"):
            update_profile(username, {"name": full_name, "email": email})
            u = load_students().get(username, {})
            st.session_state.user = {"username": username, **u}
            st.success("✅ Dados atualizados!")

        st.markdown("---")
        st.markdown("### Alterar Senha")
        col3, col4 = st.columns(2)
        with col3:
            new_pw  = st.text_input(t("new_password", ui_lang), type="password", key="pf_newpw")
        with col4:
            conf_pw = st.text_input(t("confirm_password", ui_lang), type="password", key="pf_confpw")

        if st.button(t("change_password", ui_lang), key="save_pw"):
            if len(new_pw) < 6:
                st.error("Senha muito curta.")
            elif new_pw != conf_pw:
                st.error("As senhas não coincidem.")
            else:
                update_password(username, new_pw)
                st.success("✅ Senha alterada!")

    st.markdown("---")
    back_label = "← Voltar ao Dashboard" if is_prof else "← Voltar ao Chat"
    back_page  = "dashboard" if is_prof else "chat"
    if st.button(back_label, key="back_chat"):
        st.session_state.page = back_page; st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# MODO CONVERSA — avatar animado com microfone contínuo (VAD)
# ══════════════════════════════════════════════════════════════════════════════

def _vm_process_audio(raw: bytes, lang: str, conv_id: str) -> None:
    """
    Transcreve o áudio gravado, envia ao Claude e armazena
    a resposta + TTS no session_state para o JS exibir.
    """
    txt = transcribe_bytes(raw, suffix=".webm", language=None)
    if not txt or txt.startswith("❌") or txt.startswith("⚠️"):
        st.session_state["_vm_error"] = txt or "Não entendi. Tente novamente."
        return

    st.session_state["_vm_user_said"] = txt
    if not API_KEY:
        st.session_state["_vm_error"] = "❌ ANTHROPIC_API_KEY não configurada."
        return

    user     = st.session_state.user
    username = user["username"]
    history  = st.session_state.get("_vm_history", [])
    context  = f"\n\nStudent: Name={user['name']}, Level={user['level']}, Focus={user['focus']}."

    history.append({"role": "user", "content": txt})
    client = anthropic.Anthropic(api_key=API_KEY)
    resp   = client.messages.create(
        model="claude-haiku-4-5", max_tokens=1000,
        system=SYSTEM_PROMPT + context, messages=history
    )
    reply = resp.content[0].text
    history.append({"role": "assistant", "content": reply})
    st.session_state["_vm_history"] = history

    # ── Detecta pronúncia excelente no reply ─────────────────────────────────
    _praise_en = ["great pronunciation","excellent pronunciation","perfect pronunciation",
                  "well pronounced","great accent","excellent accent","perfect accent",
                  "sounded great","sounded perfect","very clear","beautifully said",
                  "well said","that was perfect","spot on","nailed it"]
    _praise_pt = ["ótima pronúncia","excelente pronúncia","pronúncia perfeita",
                  "pronunciou muito bem","sotaque ótimo","muito claro","ficou perfeito",
                  "perfeito","mandou bem","muito bem pronunciado"]
    _reply_low = reply.lower()
    _good = any(p in _reply_low for p in _praise_en + _praise_pt)
    st.session_state["_vm_good_pronunciation"] = _good

    # Gera TTS para o modo voz
    tts_b64 = ""
    tts_bytes = None
    if tts_available():
        ab = text_to_speech(reply)
        if ab:
            tts_bytes = ab
            tts_b64 = base64.b64encode(ab).decode()

    st.session_state["_vm_reply"]   = reply
    st.session_state["_vm_tts_b64"] = tts_b64

    # Persiste no histórico da conversa atual
    append_message(username, conv_id, "user",      txt,   audio=True)
    append_message(username, conv_id, "assistant", reply, tts_b64=tts_b64 or None)


def show_voice() -> None:
    user     = st.session_state.user
    username = user["username"]
    profile  = user.get("profile", {})
    lang     = profile.get("language", "pt-BR")

    ring_color        = profile.get("ring_color",        "#f0a500")
    user_bubble_color = profile.get("user_bubble_color", "#2d6a4f")
    bot_bubble_color  = profile.get("bot_bubble_color",  "#1a1f2e")

    def _rgba(h: str, a: float) -> str:
        h = h.lstrip("#")
        if len(h) == 3: h = h[0]*2+h[1]*2+h[2]*2
        r, g, b = int(h[0:2],16), int(h[2:4],16), int(h[4:6],16)
        return f"rgba({r},{g},{b},{a})"

    # Sem scroll no modo voz
    st.markdown("""<style>
body,.stApp,[data-testid="stAppViewContainer"],[data-testid="stMain"]{background:#060a10!important;}
section[data-testid="stMain"]>div,.main .block-container{padding:0!important;margin:0!important;overflow:hidden!important;max-height:100vh!important;}
div[data-testid="stVerticalBlock"],div[data-testid="stVerticalBlockBorderWrapper"],div[data-testid="element-container"]{gap:0!important;padding:0!important;margin:0!important;}
html,body{overflow:hidden!important;}
/* Remove barra azul/header do Streamlit */
[data-testid="stHeader"],[data-testid="stDecoration"],
header[data-testid="stHeader"],div[data-testid="stDecoration"],
#MainMenu,footer,header{display:none!important;height:0!important;visibility:hidden!important;}
[data-testid="stToolbar"]{display:none!important;}
/* Garante que o app começa do topo sem espaço reservado para o header */
.stApp>[data-testid="stAppViewContainer"]{padding-top:0!important;}
[data-testid="stAppViewContainer"]{padding-top:0!important;margin-top:0!important;}
</style>""", unsafe_allow_html=True)

    conv_id = get_or_create_conv(username)

    st.markdown("""
<style>
.vm-close-btn {
    position: fixed;
    top: 14px;
    right: 16px;
    z-index: 9999;
}
.vm-close-btn button {
    background: rgba(255,255,255,0.08) !important;
    border: 1px solid rgba(255,255,255,0.15) !important;
    color: #ccc !important;
    border-radius: 8px !important;
    font-size: 0.78rem !important;
    padding: 4px 12px !important;
    cursor: pointer !important;
}
.vm-close-btn button:hover {
    background: rgba(255,255,255,0.15) !important;
    color: #fff !important;
}
</style>
""", unsafe_allow_html=True)

    with st.container():
        st.markdown('<div class="vm-close-btn">', unsafe_allow_html=True)
        if st.button("✕ Close", key="vm_close_btn"):
            st.session_state.voice_mode = False
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    # Carrega histórico do banco se _vm_history estiver vazio
    if not st.session_state.get("_vm_history") and conv_id:
        msgs_db = load_conversation(username, conv_id)
        if msgs_db:
            st.session_state["_vm_history"] = [
                {
                    "role":    m["role"],
                    "content": m["content"],
                    "tts_b64": m.get("tts_b64", ""),   # ← preserva áudio por mensagem
                }
                for m in msgs_db if m.get("content")
            ]

    # Processa áudio recebido — só 1 rerun após processar
    audio_val = st.audio_input(
        " ", key=f"voice_input_{st.session_state.audio_key}",
        label_visibility="collapsed",
    )
    if audio_val and audio_val != st.session_state.get("_vm_last_upload"):
        st.session_state["_vm_last_upload"] = audio_val
        for k in ["_vm_reply", "_vm_tts_b64", "_vm_user_said", "_vm_error"]:
            st.session_state.pop(k, None)
        with st.spinner(t("processing", lang)):
            _vm_process_audio(audio_val.read(), lang, conv_id)
        st.session_state.audio_key += 1
        st.rerun()

    # Estado atual
    reply   = st.session_state.get("_vm_reply",   "")
    tts_b64 = st.session_state.get("_vm_tts_b64", "")
    vm_error = st.session_state.get("_vm_error",  "")
    history  = st.session_state.get("_vm_history", [])

    # Frames do avatar
    frames    = get_avatar_frames()
    has_anim  = bool(frames["normal"])

    # Serializa dados para JS
    history_js  = json.dumps(history)
    tts_js      = json.dumps(tts_b64)
    reply_js    = json.dumps(reply)
    err_js      = json.dumps(vm_error)
    tap_speak   = json.dumps(t("tap_to_speak", lang))
    tap_stop    = json.dumps(t("tap_to_stop",  lang))
    speaking_   = json.dumps(t("speaking_ai",  lang))
    proc_       = json.dumps(t("processing",   lang))

    av_normal_js     = json.dumps(frames["normal"])
    av_meio_js       = json.dumps(frames["meio"])
    av_aberta_js     = json.dumps(frames["aberta"])
    av_bem_aberta_js = json.dumps(frames["bem_aberta"])
    av_ouvindo_js    = json.dumps(frames["ouvindo"])
    av_piscando_js   = json.dumps(frames["piscando"])
    av_surpresa_js   = json.dumps(frames["surpresa"])
    has_anim_js      = "true" if has_anim else "false"
    photo_js         = json.dumps(get_tati_mini_b64() or get_photo_b64())
    prof_name_js     = json.dumps(PROF_NAME)

    good_pronunc_js = json.dumps(bool(
        st.session_state.get("_vm_good_pronunciation", False)
    ))
    st.session_state.pop("_vm_good_pronunciation", None)

    components.html(f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1,viewport-fit=cover">
<style>
@import url('https://fonts.googleapis.com/css2?family=Sora:wght@300;400;600;700&display=swap');
*{{box-sizing:border-box;margin:0;padding:0;}}
html,body{{
    background:#060a10;font-family:'Sora',sans-serif;
    height:100%;overflow:hidden;
    /* iOS safe area */
    padding-bottom:env(safe-area-inset-bottom);
}}
.app{{
    display:flex;flex-direction:column;align-items:center;
    height:100vh;
    height:100dvh;
    padding:0 16px 0;
    gap:0;overflow:hidden;
}}

/* ── Avatar ── */
.avatar-section{{
    flex-shrink:0;width:100%;
    display:flex;flex-direction:column;align-items:center;gap:4px;
    padding:12px 0 8px;
    position:sticky;top:0;z-index:10;
    background:linear-gradient(180deg,#060a10 85%,transparent 100%);
}}
/* Avatar menor em telas pequenas */
.avatar-wrap{{position:relative;width:200px;height:200px;flex-shrink:0;}}
@media(max-height:700px){{
    .avatar-wrap{{width:130px;height:130px;}}
    .avatar-img,.avatar-emoji{{width:130px!important;height:130px!important;}}
    .avatar-section{{padding:6px 0 4px;}}
    .prof-name{{font-size:.85rem!important;}}
}}
.avatar-ring{{
    position:absolute;inset:-8px;border-radius:50%;
    border:2px solid {_rgba(ring_color,.3)};
    animation:ring-pulse 2s ease-in-out infinite;
}}
.avatar-ring.active{{
    border-color:{ring_color};
    box-shadow:0 0 0 0 {_rgba(ring_color,.5)};
    animation:ring-glow 1s ease-in-out infinite;
}}
@keyframes ring-pulse{{0%,100%{{opacity:.4;transform:scale(1);}}50%{{opacity:.8;transform:scale(1.03);}}}}
@keyframes ring-glow{{0%{{box-shadow:0 0 0 0 {_rgba(ring_color,.5)};}}70%{{box-shadow:0 0 14px {_rgba(ring_color,0)};}}100%{{box-shadow:0 0 0 0 {_rgba(ring_color,0)};}}}}
.avatar-img{{
    width:200px;height:200px;border-radius:50%;
    object-fit:cover;object-position:top center;
    border:3px solid {ring_color};
    box-shadow:0 0 32px {_rgba(ring_color,.25)};
}}
.avatar-emoji{{
    width:200px;height:200px;border-radius:50%;
    background:linear-gradient(135deg,#1a2535,#0f1824);
    border:3px solid {ring_color};
    display:flex;align-items:center;justify-content:center;font-size:54px;
}}
.prof-name{{font-size:1rem;font-weight:700;color:#e6edf3;margin-top:6px;}}
.status{{font-size:.68rem;color:{ring_color};margin-top:1px;}}

/* ── Histórico de bolhas ── */
.history-wrap{{
    width:100%;max-width:1890px;
    flex:1;min-height:0;
    overflow-y:auto;display:flex;flex-direction:column;gap:8px;
    padding:8px 4px;
    scrollbar-width:thin;scrollbar-color:#1a2535 transparent;
    -webkit-overflow-scrolling:touch;
}}
.history-wrap::-webkit-scrollbar{{width:4px;}}
.history-wrap::-webkit-scrollbar-thumb{{background:#1a2535;border-radius:4px;}}
.bubble{{
    max-width:82%;padding:10px 15px;border-radius:18px;
    font-size:.84rem;line-height:1.55;word-break:break-word;
}}
.bubble.user{{
    align-self:flex-end;
    background:{user_bubble_color};color:#fff;
    border-bottom-right-radius:4px;
}}
.bubble.bot{{
    align-self:flex-start;
    background:{bot_bubble_color};color:#e6edf3;
    border:1px solid {_rgba(bot_bubble_color,.8)};
    border-bottom-left-radius:4px;
}}
.bubble-label{{font-size:.6rem;color:#4a5a6a;margin:2px 4px;}}
.bubble-label.right{{text-align:right;}}

/* ── Botão de áudio por bolha ── */
.bubble-play-btn{{
    align-self:flex-start;
    background:transparent;border:1px solid #1a2535;color:#3a6a8a;
    font-size:.72rem;padding:4px 12px;border-radius:8px;
    cursor:pointer;font-family:inherit;transition:all .15s;margin-bottom:4px;
    /* touch-friendly */
    min-height:32px;
}}
.bubble-play-btn:hover{{color:#f0a500;border-color:rgba(240,165,0,.4);background:rgba(240,165,0,.06);}}
.bubble-play-btn.playing{{color:#e05c2a;border-color:rgba(224,92,42,.5);}}

/* ── Erro ── */
.error-box{{
    background:rgba(224,92,42,.1);border:1px solid rgba(224,92,42,.3);
    border-radius:10px;padding:8px 14px;font-size:.78rem;color:#e05c2a;
    max-width:560px;width:100%;text-align:center;flex-shrink:0;
}}

/* ── Rodapé do mic — RESPONSIVO ── */
.mic-footer{{
    flex-shrink:0;
    width:100%;max-width:620px;
    display:flex;flex-direction:column;align-items:center;
    gap:6px;
    padding:8px 0 max(16px, env(safe-area-inset-bottom));
    background:linear-gradient(to top,#060a10 70%,transparent);
    position:sticky;bottom:0;
}}

/* ── Controles de áudio: linha única que não quebra ── */
.audio-controls{{
    display:flex;
    align-items:center;
    gap:6px;
    padding:8px 12px;
    background:#0d1420;
    border:1px solid #1a2535;
    border-radius:12px;
    width:100%;
    overflow-x:auto;          /* scroll horizontal se necessário */
    overflow-y:hidden;
    -webkit-overflow-scrolling:touch;
    white-space:nowrap;
    scrollbar-width:none;     /* esconde scrollbar */
    flex-wrap:nowrap;         /* NUNCA quebra linha */
    min-height:44px;
}}
.audio-controls::-webkit-scrollbar{{display:none;}}

/* Em telas muito pequenas, reduz padding e fonte */
@media(max-width:400px){{
    .audio-controls{{padding:6px 10px;gap:4px;}}
    .ctrl-label{{font-size:.6rem;}}
    .ctrl-val{{font-size:.6rem;min-width:24px;}}
    #global-play-btn{{padding:4px 10px;font-size:.72rem;}}
}}

.ctrl-label{{font-size:.68rem;color:#4a5a6a;white-space:nowrap;flex-shrink:0;}}
.ctrl-val{{font-size:.68rem;color:#8b949e;min-width:28px;text-align:left;flex-shrink:0;}}

input[type=range].ctrl-range{{
    -webkit-appearance:none;
    flex-shrink:0;
    width:60px;
    height:4px;
    background:#1a2535;border-radius:2px;outline:none;cursor:pointer;
    touch-action:none;   /* evita scroll ao arrastar no mobile */
}}
@media(min-width:480px){{
    input[type=range].ctrl-range{{ width:80px; }}
}}
input[type=range].ctrl-range::-webkit-slider-thumb{{
    -webkit-appearance:none;width:16px;height:16px;  /* maior para touch */
    border-radius:50%;background:{ring_color};cursor:pointer;
}}
input[type=range].ctrl-range::-moz-range-thumb{{
    width:16px;height:16px;border-radius:50%;background:{ring_color};cursor:pointer;border:none;
}}

#global-play-btn{{
    background:#1a2535;color:#e6edf3;border:1px solid #252d3d;
    border-radius:8px;padding:5px 12px;font-size:.78rem;cursor:pointer;
    white-space:nowrap;transition:background .15s;font-family:inherit;
    flex-shrink:0;
    min-height:32px;        /* touch-friendly */
    touch-action:manipulation;
}}
#global-play-btn:hover{{background:#252d3d;}}

/* ── Botão mic ── */
.mic-btn{{
    width:72px;height:72px;border-radius:50%;border:none;cursor:pointer;
    background:linear-gradient(135deg,#1a2535,#131c2a);
    color:#8b949e;font-size:28px;
    display:flex;align-items:center;justify-content:center;
    box-shadow:0 4px 20px rgba(0,0,0,.4),inset 0 1px 0 rgba(255,255,255,.05);
    transition:all .2s;outline:none;
    touch-action:manipulation;   /* remove delay 300ms no mobile */
    -webkit-tap-highlight-color:transparent;
    flex-shrink:0;
}}
/* Mic maior em tablets */
@media(min-width:600px){{
    .mic-btn{{width:80px;height:80px;font-size:32px;}}
}}
.mic-btn:hover{{background:linear-gradient(135deg,#1e2f40,#182130);color:#e6edf3;}}
.mic-btn.recording{{
    background:linear-gradient(135deg,#e05c2a,#c44a1a);color:#fff;
    box-shadow:0 0 0 0 rgba(224,92,42,.6),0 4px 20px rgba(224,92,42,.3);
    animation:mic-pulse 1.2s ease-in-out infinite;
}}
.mic-btn.processing{{
    background:linear-gradient(135deg,#f0a500,#c88800);color:#060a10;animation:none;
}}
@keyframes mic-pulse{{
    0%{{box-shadow:0 0 0 0 rgba(224,92,42,.6),0 4px 20px rgba(224,92,42,.3);}}
    70%{{box-shadow:0 0 0 16px rgba(224,92,42,0),0 4px 20px rgba(224,92,42,.3);}}
    100%{{box-shadow:0 0 0 0 rgba(224,92,42,0),0 4px 20px rgba(224,92,42,.3);}}
}}
.mic-hint{{font-size:.68rem;color:#4a5a6a;letter-spacing:.3px;}}
</style>
</head><body>
<div class="app" id="app">
    <div class="avatar-section">
        <div class="avatar-wrap">
            <div class="avatar-ring" id="ring"></div>
            <img id="avImg" class="avatar-img" src="" alt="" style="display:none;"
                 onerror="this.style.display='none';document.getElementById('avEmoji').style.display='flex';">
            <div id="avEmoji" class="avatar-emoji">&#129489;&#8205;&#127979;</div>
        </div>
        <div class="prof-name" id="profName"></div>
        <div class="status" id="statusTxt">&#9679; Online</div>
    </div>

    <div class="history-wrap" id="historyWrap"></div>
    <div class="error-box" id="errBox" style="display:none;"></div>

    <div class="mic-footer">
        <div class="audio-controls" id="audioControls">
            <button id="global-play-btn">&#9654; Ouvir</button>
            <span class="ctrl-label">Vol</span>
            <input type="range" class="ctrl-range" id="vol-slider" min="0" max="1" step="0.05" value="1">
            <span class="ctrl-val" id="vol-val">100%</span>
            <span class="ctrl-label">Vel</span>
            <input type="range" class="ctrl-range" id="spd-slider" min="0.5" max="2" step="0.1" value="1">
            <span class="ctrl-val" id="spd-val">1.0x</span>
        </div>
        <button class="mic-btn" id="micBtn"><i class="fa-solid fa-microphone"></i></button>
        <div class="mic-hint" id="micHint"></div>
    </div>
</div>

<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.1/css/all.min.css">
<script>
(function(){{
// ══════════════════════════════════════════════════════════════════════════════
// DADOS DO PYTHON
// ══════════════════════════════════════════════════════════════════════════════
var TTS_B64      = {tts_js};
var REPLY        = {reply_js};
var HISTORY      = {history_js};
var VM_ERROR     = {err_js};
var TAP_SPEAK    = {tap_speak};
var TAP_STOP     = {tap_stop};
var SPEAKING     = {speaking_};
var HAS_ANIM     = {has_anim_js};
var GOOD_PRONUNC = {good_pronunc_js};
var PHOTO        = {photo_js};
var PROF_NAME    = {prof_name_js};

// 7 frames
var F_NORMAL     = {av_normal_js};
var F_MEIO       = {av_meio_js};
var F_ABERTA     = {av_aberta_js};
var F_BEM_ABERTA = {av_bem_aberta_js};
var F_OUVINDO    = {av_ouvindo_js};
var F_PISCANDO   = {av_piscando_js};
var F_SURPRESA   = {av_surpresa_js};

// ══════════════════════════════════════════════════════════════════════════════
// ELEMENTOS
// ══════════════════════════════════════════════════════════════════════════════
var micBtn    = document.getElementById('micBtn');
var micHint   = document.getElementById('micHint');
var statusTxt = document.getElementById('statusTxt');
var errBox    = document.getElementById('errBox');
var ring      = document.getElementById('ring');
var avImg     = document.getElementById('avImg');
var avEmoji   = document.getElementById('avEmoji');
var histWrap  = document.getElementById('historyWrap');
var profName  = document.getElementById('profName');

profName.textContent = PROF_NAME;
micHint.textContent  = TAP_SPEAK;

// ══════════════════════════════════════════════════════════════════════════════
// CONTROLE DE FRAME
// ══════════════════════════════════════════════════════════════════════════════
var _lastFrame = '';
function setFrame(src){{
    if(!src || src === _lastFrame) return;
    _lastFrame = src;
    avImg.src  = src;
    avImg.style.display   = 'block';
    avEmoji.style.display = 'none';
}}
setFrame(HAS_ANIM ? F_NORMAL : (PHOTO || F_NORMAL));

// ══════════════════════════════════════════════════════════════════════════════
// MÁQUINA DE ESTADOS:  idle | listening | processing | speaking
// ══════════════════════════════════════════════════════════════════════════════
var _state      = 'idle';
var _blinkTimer = null;
var _mouthTimer = null;
var _analyser   = null;
var _audioCtx   = null;
var _mouthFallbackIdx = 0;

function _stopAllTimers(){{
    if(_blinkTimer){{ clearTimeout(_blinkTimer); clearInterval(_blinkTimer); _blinkTimer = null; }}
    if(_mouthTimer){{ clearInterval(_mouthTimer); _mouthTimer = null; }}
}}

// ── IDLE: normal + piscar natural (3-5s, 150ms fechado) ──────────────────────
function enterIdle(){{
    _stopAllTimers();
    _state = 'idle';
    setFrame(F_NORMAL);
    ring.classList.remove('active');
    statusTxt.textContent = '● Online';
    function scheduleBlink(){{
        var delay = 3210 + Math.random() * 2000;
        _blinkTimer = setTimeout(function(){{
            if(_state !== 'idle') return;
            setFrame(F_PISCANDO);
            setTimeout(function(){{
                if(_state !== 'idle') return;
                setFrame(F_NORMAL);
                scheduleBlink();
            }}, 150);
        }}, delay);
    }}
    scheduleBlink();
}}

// ── LISTENING: ouvindo fixo enquanto mic está ativo ──────────────────────────
function enterListening(){{
    _stopAllTimers();
    _state = 'listening';
    setFrame(F_OUVINDO);
    ring.classList.remove('active');
    statusTxt.textContent = '🎙 Ouvindo…';
}}

// ── PROCESSING: normal com piscada lenta enquanto Claude pensa ───────────────
function enterProcessing(){{
    _stopAllTimers();
    _state = 'processing';
    setFrame(F_NORMAL);
    ring.classList.remove('active');
    statusTxt.textContent = '⏳ Processando…';
    _blinkTimer = setInterval(function(){{
        if(_state !== 'processing') return;
        setFrame(F_PISCANDO);
        setTimeout(function(){{
            if(_state !== 'processing') return;
            setFrame(F_NORMAL);
        }}, 180);
    }}, 2200);
}}

// ── SPEAKING: sincronização labial via Web Audio API ─────────────────────────
// normal = boca fechada (pausas) | meio = boca aberta (fala)
function enterSpeaking(audioEl){{
    _stopAllTimers();
    _state = 'speaking';
    ring.classList.add('active');
    statusTxt.textContent = SPEAKING;

    if(!HAS_ANIM){{ return; }}

    try{{
        if(!_audioCtx){{
            _audioCtx = new (window.AudioContext || window.webkitAudioContext)();
        }}
        if(!_analyser){{
            _analyser = _audioCtx.createAnalyser();
            _analyser.fftSize = 1024;
            _analyser.smoothingTimeConstant = 0.1; // reage rápido ao volume real
            var src = _audioCtx.createMediaElementSource(audioEl);
            src.connect(_analyser);
            _analyser.connect(_audioCtx.destination);
        }}

        var buf = new Uint8Array(_analyser.frequencyBinCount);
        _mouthTimer = setInterval(function(){{
            if(_state !== 'speaking') return;
            _analyser.getByteFrequencyData(buf);
            // média simples dos bins de fala
            var sum = 0, n = Math.min(100, buf.length);
            for(var i = 4; i < n; i++) sum += buf[i];
            var avg = sum / (n - 4); // 0..255
            setFrame(avg < 18 ? F_NORMAL : F_MEIO);
        }}, 60); // ~16fps — rápido o suficiente para parecer natural

    }}catch(e){{
        // Fallback sem Web Audio: alterna meio↔normal a ~2Hz
        _mouthFallbackIdx = 0;
        _mouthTimer = setInterval(function(){{
            if(_state !== 'speaking') return;
            setFrame(_mouthFallbackIdx % 2 === 0 ? F_MEIO : F_NORMAL);
            _mouthFallbackIdx++;
        }}, 250);
    }}
}}

// ── Fim da fala ───────────────────────────────────────────────────────────────
function onSpeakingEnded(goodPronunc){{
    _stopAllTimers();
    if(goodPronunc && F_BEM_ABERTA){{
        // Pronúncia excelente: reação de aprovação por 1.2s
        setFrame(F_BEM_ABERTA);
        setTimeout(function(){{ enterIdle(); }}, 1200);
    }} else {{
        enterIdle();
    }}
}}

// ══════════════════════════════════════════════════════════════════════════════
// ÁUDIO GLOBAL
// ══════════════════════════════════════════════════════════════════════════════
var currentAudio = null, lastB64 = null;
function getVol(){{ return parseFloat(document.getElementById('vol-slider').value) || 1; }}
function getSpd(){{ return parseFloat(document.getElementById('spd-slider').value) || 1; }}

function playTTS(b64, onEndCallback){{
    if(currentAudio){{ currentAudio.pause(); currentAudio = null; }}
    _analyser = null;
    if(!b64) return;
    lastB64 = b64;
    var audio = new Audio('data:audio/mp3;base64,' + b64);
    audio.volume = getVol(); audio.playbackRate = getSpd(); audio._srcB64 = b64;
    currentAudio = audio;
    audio.onplay   = function(){{ enterSpeaking(audio); updateGlobalBtn(true); }};
    audio.onended  = function(){{
        currentAudio = null; updateGlobalBtn(false);
        onSpeakingEnded(GOOD_PRONUNC);
        if(onEndCallback) onEndCallback();
    }};
    audio.onerror  = function(){{ currentAudio = null; updateGlobalBtn(false); enterIdle(); }};
    audio.play().catch(function(){{ currentAudio = null; updateGlobalBtn(false); enterIdle(); }});
}}
function stopTTS(){{
    if(currentAudio){{ currentAudio.pause(); currentAudio = null; }}
    _analyser = null;
    updateGlobalBtn(false); enterIdle();
}}
function updateGlobalBtn(playing){{
    var btn = document.getElementById('global-play-btn');
    if(!btn) return;
    btn.textContent  = playing ? '⏹ Parar' : '▶ Ouvir';
    btn.style.background = playing ? '#8b2a2a' : '#1a2535';
}}

document.getElementById('global-play-btn').addEventListener('click', function(){{
    if(currentAudio && !currentAudio.paused) stopTTS();
    else if(lastB64 || TTS_B64) playTTS(lastB64 || TTS_B64);
}});
document.getElementById('vol-slider').addEventListener('input', function(){{
    document.getElementById('vol-val').textContent = Math.round(this.value * 100) + '%';
    if(currentAudio) currentAudio.volume = parseFloat(this.value);
}});
document.getElementById('spd-slider').addEventListener('input', function(){{
    document.getElementById('spd-val').textContent = parseFloat(this.value).toFixed(1) + 'x';
    if(currentAudio) currentAudio.playbackRate = parseFloat(this.value);
}});

// ══════════════════════════════════════════════════════════════════════════════
// BOLHAS DE HISTÓRICO
// ══════════════════════════════════════════════════════════════════════════════
function addBubble(role, text, b64){{
    var label = document.createElement('div');
    label.className = 'bubble-label' + (role === 'user' ? ' right' : '');
    label.textContent = role === 'user' ? 'Você' : PROF_NAME;
    var bub = document.createElement('div');
    bub.className = 'bubble ' + role;
    bub.textContent = text;
    histWrap.appendChild(label);
    histWrap.appendChild(bub);
    if(role === 'bot' && b64){{
        var pbtn = document.createElement('button');
        pbtn.className = 'bubble-play-btn';
        pbtn.textContent = '▶ Ouvir';
        pbtn.addEventListener('click', function(){{
            var isPlaying = currentAudio && !currentAudio.paused && currentAudio._srcB64 === b64;
            if(isPlaying){{ stopTTS(); pbtn.textContent = '▶ Ouvir'; pbtn.classList.remove('playing'); }}
            else{{
                document.querySelectorAll('.bubble-play-btn').forEach(function(b){{
                    b.textContent = '▶ Ouvir'; b.classList.remove('playing');
                }});
                pbtn.textContent = '⏹ Parar'; pbtn.classList.add('playing');
                playTTS(b64, function(){{ pbtn.textContent = '▶ Ouvir'; pbtn.classList.remove('playing'); }});
            }}
        }});
        histWrap.appendChild(pbtn);
    }}
    histWrap.scrollTop = histWrap.scrollHeight;
}}

// ══════════════════════════════════════════════════════════════════════════════
// RENDERIZA ESTADO INICIAL
// ══════════════════════════════════════════════════════════════════════════════
if(VM_ERROR){{
    errBox.textContent = VM_ERROR; errBox.style.display = 'block'; enterIdle();
}} else {{
    errBox.style.display = 'none';
    if(HISTORY && HISTORY.length > 0){{
        HISTORY.forEach(function(msg){{
            addBubble(msg.role === 'user' ? 'user' : 'bot', msg.content, msg.tts_b64 || '');
        }});
    }}
    if(TTS_B64) setTimeout(function(){{ playTTS(TTS_B64); }}, 300);
    else        enterIdle();
}}

// ══════════════════════════════════════════════════════════════════════════════
// MICROFONE
// ══════════════════════════════════════════════════════════════════════════════
var recording = false;
function getRealMicBtn(){{
    var doc = window.parent.document;
    var ai  = doc.querySelector('[data-testid="stAudioInput"]');
    if(!ai) return null;
    return ai.querySelector('button') || ai.querySelector('[data-testid="stAudioInputRecordButton"]');
}}
micBtn.addEventListener('click', function(){{
    var realBtn = getRealMicBtn();
    if(!realBtn) return;
    if(recording){{
        recording = false;
        micBtn.classList.remove('recording');
        micBtn.classList.add('processing');
        micBtn.innerHTML = '<i class="fa-solid fa-spinner fa-spin"></i>';
        micHint.textContent = TAP_SPEAK;
        enterProcessing();
        realBtn.click();
    }} else {{
        if(currentAudio){{ currentAudio.pause(); currentAudio = null; }}
        if(window.parent.speechSynthesis) window.parent.speechSynthesis.cancel();
        recording = true;
        micBtn.classList.remove('processing');
        micBtn.classList.add('recording');
        micBtn.innerHTML = '<i class="fa-solid fa-stop"></i>';
        micHint.textContent = TAP_STOP;
        enterListening();
        realBtn.click();
    }}
}});

// ── Esconde stAudioInput nativo ──
function hideNativeAudio(){{
    var doc=window.parent.document;
    var ai=doc.querySelector('[data-testid="stAudioInput"]');
    if(ai){{
        ai.style.cssText='position:fixed;bottom:-999px;left:-9999px;opacity:0;pointer-events:none;width:1px;height:1px;';
        var btn=ai.querySelector('button');
        if(btn) btn.style.pointerEvents='auto';
    }}
}}
hideNativeAudio();
try{{
    var obs=new MutationObserver(hideNativeAudio);
    obs.observe(window.parent.document.body,{{childList:true,subtree:true}});
    setTimeout(function(){{obs.disconnect();}},15000);
}}catch(e){{}}

// ── Resize iframe — usa dvh para mobile (esconde barra do browser) ──
(function resizeIframe(){{
    try{{
        var par = window.parent;
        // Altura real do viewport (dvh = dynamic viewport height, funciona no mobile)
        var h = par.innerHeight;
        try{{
            // visualViewport é mais preciso no mobile (exclui teclado virtual)
            if(par.visualViewport) h = par.visualViewport.height;
        }}catch(e){{}}

        var iframes = par.document.querySelectorAll('iframe');
        for(var i=0;i<iframes.length;i++){{
            try{{
                if(iframes[i].contentWindow===window){{
                    iframes[i].style.cssText=[
                        'height:'+h+'px',
                        'max-height:'+h+'px',
                        'min-height:200px',
                        'display:block',
                        'border:none',
                        'width:100%',
                    ].join(';');
                    // Remove padding/margin dos wrappers Streamlit
                    var p=iframes[i].parentElement;
                    for(var j=0;j<10&&p&&p!==par.document.body;j++){{
                        p.style.margin='0';p.style.padding='0';
                        p.style.overflow='hidden';p.style.maxHeight=h+'px';
                        p=p.parentElement;
                    }}
                    break;
                }}
            }}catch(e){{}}
        }}
    }}catch(e){{}}

    // Re-executa ao redimensionar E ao mudar visualViewport (teclado mobile)
    try{{
        par.removeEventListener('resize',resizeIframe);
        par.addEventListener('resize',resizeIframe);
        if(par.visualViewport){{
            par.visualViewport.removeEventListener('resize',resizeIframe);
            par.visualViewport.addEventListener('resize',resizeIframe);
        }}
    }}catch(e){{}}
}})();

}})();
</script>
</body></html>""", height=920, scrolling=False)


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DO CHAT — processamento de arquivos
# ══════════════════════════════════════════════════════════════════════════════

def _process_and_send_file(username: str, user: dict, conv_id: str,
                            raw: bytes, filename: str, extra_text: str = "") -> bool:
    """
    Processa um arquivo (áudio / texto / imagem) e envia ao Claude
    com um texto adicional opcional do usuário.
    Retorna True se o envio foi bem-sucedido.
    """
    result = extract_file(raw, filename)
    kind, label = result["kind"], result["label"]

    if kind == "audio":
        with st.spinner("🔄 Transcrevendo áudio..."):
            text = transcribe_bytes(raw, suffix=Path(filename).suffix.lower(), language="en")
        if text.startswith("❌") or text.startswith("⚠️"):
            st.error(text); return False
        user_display = f"{extra_text}\n\n[Áudio transcrito: {text}]" if extra_text else text
        claude_msg   = f"{extra_text}\n\n[Áudio: '{filename}']\n{text}" if extra_text else \
                       f"[Áudio: '{filename}']\n{text}"
        append_message(username, conv_id, "user", user_display, audio=True)
        st.session_state.speaking = True
        try:   send_to_claude(username, user, conv_id, claude_msg)
        except Exception as e: st.error(f"❌ {e}")
        st.session_state.speaking = False
        return True

    elif kind == "text":
        extracted = result["text"]
        if extracted.startswith("❌"):  st.error(extracted);  return False
        if not extracted:               st.warning(f"Sem texto em '{filename}'."); return False
        preview      = extracted[:200].replace('\n', ' ')
        user_display = f"📄 [{label}: '{filename}'] — {preview}{'…' if len(extracted)>200 else ''}"
        if extra_text: user_display = f"{extra_text}\n\n{user_display}"
        claude_msg   = (f"📄 [{label}: '{filename}']\n\n{extracted}\n\n"
                        "Please help me understand this content — explain vocabulary, grammar, and key ideas.")
        if extra_text: claude_msg = f"{extra_text}\n\n{claude_msg}"
        append_message(username, conv_id, "user", user_display)
        st.session_state.speaking = True
        try:   send_to_claude(username, user, conv_id, claude_msg)
        except Exception as e: st.error(f"❌ {e}")
        st.session_state.speaking = False
        return True

    elif kind == "image":
        user_display = f"📸 [Imagem: '{filename}']"
        if extra_text: user_display = f"{extra_text}\n\n{user_display}"
        claude_msg   = f"📸 [Imagem: '{filename}']\nPlease look at this image and help me learn English from it."
        if extra_text: claude_msg = f"{extra_text}\n\n{claude_msg}"
        append_message(username, conv_id, "user", user_display)
        st.session_state.speaking = True
        try:
            send_to_claude(username, user, conv_id, claude_msg,
                           image_b64=result["b64"], image_media_type=result["media_type"])
        except Exception as e: st.error(f"❌ {e}")
        st.session_state.speaking = False
        return True

    else:
        st.warning(f"⚠️ Formato '{label}' não suportado.")
        return False


# ══════════════════════════════════════════════════════════════════════════════
# CHAT — tela principal de conversa
# ══════════════════════════════════════════════════════════════════════════════

def _logout() -> None:
    """
    Encerra a sessão: apaga o token do banco SQLite e do localStorage,
    e limpa o session_state.
    """
    token = st.session_state.get("_session_token", "")
    if token:
        delete_session(token)          # remove do banco SQLite
    js_clear_session()                 # remove do localStorage do browser
    st.session_state.pop("_session_token", None)
    st.session_state.pop("_session_saved", None)
    st.session_state.update(logged_in=False, user=None, conv_id=None)


def show_chat() -> None:
    """Tela principal do chat — sidebar com histórico + área de mensagens."""
    user     = st.session_state.user
    username = user["username"]
    profile  = user.get("profile", {})
    ui_lang  = profile.get("language", "pt-BR")
    conv_id  = get_or_create_conv(username)
    messages = cached_load_conversation(username, conv_id)
    speaking = st.session_state.speaking

    # Redireciona para o modo voz se ativo
    if st.session_state.voice_mode:
        show_voice()
        return

    # ── Injeta cor de destaque do usuário no Streamlit principal ────────
    _ac = profile.get("accent_color", "#f0a500")
    _ub = profile.get("user_bubble_color", "#2d6a4f")
    _ab = profile.get("ai_bubble_color", "#1a1f2e")
    components.html(f"""<!DOCTYPE html><html><head>
<style>html,body{{margin:0;padding:0;overflow:hidden;}}</style>
</head><body><script>
(function(){{
  // ── Cores de destaque ──
  function hexToRgb(h){{
    h=h.replace('#','');
    if(h.length===3) h=h[0]+h[0]+h[1]+h[1]+h[2]+h[2];
    var n=parseInt(h,16);
    return [(n>>16)&255,(n>>8)&255,n&255].join(',');
  }}
  function luminance(h){{
    h=h.replace('#','');
    if(h.length===3) h=h[0]+h[0]+h[1]+h[1]+h[2]+h[2];
    var n=parseInt(h,16);
    var r=(n>>16)&255, g=(n>>8)&255, b=n&255;
    return 0.299*r + 0.587*g + 0.114*b;
  }}
  var ac="{_ac}", ub="{_ub}", ab="{_ab}";
  var rgb=hexToRgb(ac);
  var r=window.parent.document.documentElement;
  r.style.setProperty('--accent-full',ac);
  r.style.setProperty('--accent-70','rgba('+rgb+',.7)');
  r.style.setProperty('--accent-40','rgba('+rgb+',.4)');
  r.style.setProperty('--accent-30','rgba('+rgb+',.3)');
  r.style.setProperty('--accent-15','rgba('+rgb+',.15)');
  r.style.setProperty('--bubble-bg','rgba('+rgb+',.12)');
  r.style.setProperty('--bubble-border','rgba('+rgb+',.3)');
  r.style.setProperty('--bubble-text','#e6edf3');
  r.style.setProperty('--user-bubble-bg', ub);
  r.style.setProperty('--user-bubble-text', luminance(ub)>128 ? '#111' : '#e6edf3');
  r.style.setProperty('--ai-bubble-bg', ab);
  r.style.setProperty('--ai-bubble-text', luminance(ab)>128 ? '#111' : '#e6edf3');
  r.style.setProperty('--ai-bubble-border', 'rgba('+hexToRgb(ab)+', .6)');
  // ── Para todo áudio ao trocar de conversa ──
  var par=window.parent;
  if(par){{
    par.document.querySelectorAll('audio').forEach(function(a){{a.pause();a.currentTime=0;}});
    par.document.querySelectorAll('iframe').forEach(function(f){{
      try{{
        f.contentDocument.querySelectorAll('audio').forEach(function(a){{a.pause();a.currentTime=0;}});
        f.contentDocument.querySelectorAll('#b').forEach(function(b){{b.textContent='\u25b6 Ouvir';}});
        if(f.contentWindow.speechSynthesis) f.contentWindow.speechSynthesis.cancel();
      }}catch(e){{}}
    }});
    if(par.speechSynthesis) par.speechSynthesis.cancel();
  }}
}})();
</script></body></html>""", height=1)

    # ── Sidebar ───────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown("""<style>
        section[data-testid="stSidebar"] { overflow: hidden; }
        section[data-testid="stSidebar"] > div:first-child {
            height: 100vh; display: flex; flex-direction: column;
            padding: 0 !important; gap: 0;
        }
        div.sidebar-footer { margin-top: auto; }
        </style>""", unsafe_allow_html=True)

        # Topo com avatar da professora
        st.markdown(f"""<div style="padding:14px 14px 10px;border-bottom:1px solid #21262d;flex-shrink:0;">
            <div style="display:flex;align-items:center;gap:10px;">
                {avatar_html(40)}<div>
                <div style="font-weight:600;font-size:.88rem;">{PROF_NAME}</div>
                <div style="font-size:.68rem;color:#8b949e;"><span class="status-dot"></span>Online</div>
                </div></div></div>""", unsafe_allow_html=True)

        if st.button(t("new_conv", ui_lang), use_container_width=True, key="btn_new"):
            st.session_state.conv_id = new_conversation(username); st.rerun()
        if st.button(t("voice_mode", ui_lang), use_container_width=True, key="btn_voice"):
            st.session_state.voice_mode = True; st.rerun()

        st.markdown('<div style="font-size:.68rem;color:#8b949e;text-transform:uppercase;'
                    'letter-spacing:1px;padding:10px 4px 4px;">Conversas</div>',
                    unsafe_allow_html=True)

        convs = list_conversations(username)
        if not convs:
            st.markdown('<div style="font-size:.78rem;color:#8b949e;padding:6px 4px;">Nenhuma conversa ainda.</div>',
                        unsafe_allow_html=True)
        for c in convs:
            is_active = c["id"] == conv_id
            label     = ("▶ " if is_active else "") + c["title"]
            col_conv, col_del = st.columns([5, 1])
            with col_conv:
                if st.button(label, key=f"conv_{c['id']}", use_container_width=True,
                             help=f"📅 {c['date']} · 💬 {c['count']} msgs"):
                    st.session_state.conv_id = c["id"]; st.rerun()
            with col_del:
                if st.button("🗑", key=f"del_{c['id']}", help=t("delete_conv", ui_lang)):
                    delete_conversation(username, c["id"])
                    if st.session_state.conv_id == c["id"]:
                        st.session_state.conv_id = None
                    st.rerun()
            st.markdown(f'<div style="font-size:.62rem;color:#6e7681;margin:-10px 0 2px 6px;">'
                        f'📅 {c["date"]} · 💬 {c["count"]} msg</div>', unsafe_allow_html=True)

        # ── Rodapé da sidebar com botões de ação ──────────────────────────────
        user_msgs   = len([m for m in messages if m["role"] == "user"])
        uav_sidebar = user_avatar_html(username, size=34, fallback_emoji="🎓")
        st.markdown('<div class="sidebar-footer">', unsafe_allow_html=True)
        st.markdown("<hr style='border-color:#21262d;margin:8px 0 0'>", unsafe_allow_html=True)
        st.markdown(f"""<div style="padding:8px 12px;display:flex;align-items:center;gap:10px;">
            {uav_sidebar}
            <div style="flex:1;min-width:0;">
              <div style="font-weight:600;font-size:.82rem;white-space:nowrap;
                overflow:hidden;text-overflow:ellipsis;">{user['name'].split()[0]}</div>
              <div style="color:#8b949e;font-size:.68rem;">{user['level']} · {user_msgs} msgs</div>
            </div></div>""", unsafe_allow_html=True)

        if user["role"] == "professor":
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button(t("dashboard", ui_lang), use_container_width=True, key="btn_dash"):
                    st.session_state.page = "dashboard"; st.rerun()
            with col_b:
                if st.button(t("profile", ui_lang), use_container_width=True, key="btn_profile"):
                    st.session_state.page = "profile"; st.rerun()
            # Logout da professora: apaga sessão persistente
            if st.button(t("logout", ui_lang), use_container_width=True, key="btn_sair"):
                _logout(); st.rerun()
        else:
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button(t("profile", ui_lang), use_container_width=True, key="btn_profile"):
                    st.session_state.page = "profile"; st.rerun()
            with col_b:
                # Logout do aluno: apaga sessão persistente
                if st.button(t("logout", ui_lang), use_container_width=True, key="btn_sair"):
                    _logout(); st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    # ── CSS do chat estilo ChatGPT ────────────────────────────────────────────
    st.markdown("""<style>
[data-testid="stChatInput"] textarea {
    max-height: 120px !important; min-height: 44px !important; font-size: .88rem !important;
}
[data-testid="stChatInputContainer"] { padding: 6px 10px !important; }
.main .block-container { padding-bottom: 80px !important; }
section[data-testid="stMain"] { transition: margin-left 0.3s, width 0.3s ease !important; }

/* Mensagens estilo ChatGPT */
.msg-row { display:flex; align-items:flex-end; gap:10px; margin:6px 0; }
.msg-row.user-row { flex-direction:row-reverse; justify-content:flex-start; }
.msg-row.bot-row  { flex-direction:row; }
.msg-row.user-row > div { display:flex; flex-direction:column; align-items:flex-end; }
.msg-row.bot-row > div  { display:flex; flex-direction:column; align-items:flex-start; }

.msg-bubble {
    padding: 10px 15px; border-radius: 18px;
    font-size: .88rem; line-height: 1.6;
    word-break: normal;
    overflow-wrap: break-word;
    white-space: pre-wrap;
}
.msg-bubble.user {
    max-width: 75%;
    background: var(--user-bubble-bg, #2d6a4f);
    color: var(--user-bubble-text, #d8f3dc);
    border-bottom-right-radius: 4px;
}
.msg-bubble.bot {
    max-width: 75%;
    background: var(--ai-bubble-bg, #1a1f2e);
    color: var(--ai-bubble-text, #e6edf3);
    border: 1px solid var(--ai-bubble-border, #252d3d);
    border-bottom-left-radius: 4px;
}
.msg-bubble.audio-msg { font-style: italic; opacity: .85; }

.msg-av {
    width: 30px; height: 30px; border-radius: 50%;
    overflow: hidden; flex-shrink: 0; margin-bottom: 2px;
}
.msg-av img { width:100%; height:100%; object-fit:cover; object-position:top; }
.msg-av .av-emoji {
    width:100%; height:100%; background:#1e2a3a;
    display:flex; align-items:center; justify-content:center; font-size:14px;
}
.msg-time {
    font-size: .6rem; color: #4a5a6a;
    margin: 2px 4px 0; text-align: right;
}
.bot-row .msg-time { text-align: left; }

/* Botão Ouvir inline — sem iframe */
.msg-ouvir-row { padding: 2px 0 0 40px; }
.msg-ouvir-btn {
    background: none; border: 1px solid #30363d; border-radius: 16px;
    color: #8b949e; font-size: .68rem; padding: 2px 10px; cursor: pointer;
    transition: all .15s; white-space: nowrap; font-family: inherit;
}
.msg-ouvir-btn:hover, .msg-ouvir-btn.speaking {
    border-color: #f0a500; color: #f0a500;
}

@media (max-width: 768px) {
    .msg-bubble { max-width: 88% !important; font-size: .82rem !important; }
}
@media (max-width: 480px) {
    .msg-bubble { max-width: 94% !important; }
}
</style>""", unsafe_allow_html=True)

    # ── Header do chat ────────────────────────────────────────────────────────
    st.markdown(f"""<div class="prof-header">
        {avatar_html(56, speaking)}
        <div class="prof-info">
            <h1>{PROF_NAME}</h1>
            <p><span class="status-dot"></span>Online · {user['level']} · {user['focus']}</p>
        </div></div>""", unsafe_allow_html=True)

    # ── Histórico de mensagens ────────────────────────────────────────────────
    # Mini-avatar da Tati — background-image no div (sem <img>, sem flash)
    _tati_mini   = get_tati_mini_b64()
    tati_av_html = (f'<div class="msg-av" style="background:url({_tati_mini}) center top/cover no-repeat;"></div>'
                    if _tati_mini else
                    '<div class="msg-av"><div class="av-emoji">🧑‍🏫</div></div>')

    st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
    for i, msg in enumerate(messages):
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

        if msg["role"] == "assistant":
            is_file = msg.get("is_file", False)
            lang = st.session_state.user.get("profile", {}).get("language", "pt-BR")
            is_en = "en" in lang.lower()
            lbl_play = "Listen" if is_en else "Ouvir"
            lbl_stop = "Pause" if is_en else "Pausar"
            
            # --- SE A MENSAGEM TEM ÁUDIO B64 SALVO NO BANCO ---
            if msg.get("tts_b64"):
                b64_data = msg["tts_b64"]
                audio_html = f"""
                <div class="msg-ouvir-row" style="margin-top: 4px; margin-bottom: 8px;">
                    <audio id="audio-app-{i}" src="data:audio/mp3;base64,{b64_data}" style="display:none;"></audio>
                    <button onclick="
                        var btn = this;
                        var aud = document.getElementById('audio-app-{i}');
                        
                        if(aud.paused) {{ 
                            aud.play(); 
                            btn.innerHTML = '⏹ {lbl_stop}'; 
                            btn.style.color = '#e05c2a';
                            btn.style.borderColor = 'rgba(224,92,42,.5)';
                        }} else {{ 
                            aud.pause(); 
                            aud.currentTime = 0;
                            btn.innerHTML = '▶ {lbl_play}'; 
                            btn.style.color = '#3a6a8a';
                            btn.style.borderColor = '#1a2535';
                        }}
                        aud.onended = () => {{ 
                            btn.innerHTML = '▶ {lbl_play}'; 
                            btn.style.color = '#3a6a8a';
                            btn.style.borderColor = '#1a2535';
                        }};
                    " style="
                        background: transparent; border: 1px solid #1a2535; 
                        color: #3a6a8a; border-radius: 8px; 
                        padding: 4px 12px; cursor: pointer; font-size: 0.75rem;
                        transition: all 0.15s; font-family: sans-serif;
                    ">
                        ▶ {lbl_play}
                    </button>
                </div>
                """
                st.markdown(audio_html, unsafe_allow_html=True)
                
            # --- SE NÃO TEM B64, MAS NÃO É ARQUIVO (Usa o TTS antigo do navegador) ---
            elif not is_file:

                # Limpa o texto para o leitor de voz do navegador (fallback)
                clean_text = (msg["content"]
                    .replace("\\", "\\\\").replace("`", "")
                    .replace('"', '&quot;').replace("'", "&#39;")
                    .replace("\n", " ").replace("\r", "")
                    .replace("*", "").replace("#", ""))[:600]
                
                has_b64 = bool(msg.get("tts_b64"))
                b64_data = msg.get("tts_b64", "")
                
                # 2. Renderiza o botão no estilo do voice.py com a lógica universal
                audio_html = f"""
                <div class="msg-ouvir-row" style="margin-top: 4px; margin-bottom: 8px;">
                    <audio id="audio-app-{i}" src="data:audio/mp3;base64,{b64_data}" style="display:none;"></audio>
                    <button onclick="
                        var btn = this;
                        var aud = document.getElementById('audio-app-{i}');
                        var has_b64 = {'true' if has_b64 else 'false'};
                        
                        if(has_b64) {{
                            // Se tem áudio gerado pelo Claude/API, toca ele
                            if(aud.paused) {{ 
                                aud.play(); 
                                btn.innerHTML = '⏹ {lbl_stop}'; 
                                btn.style.color = '#e05c2a';
                                btn.style.borderColor = 'rgba(224,92,42,.5)';
                            }} else {{ 
                                aud.pause(); 
                                aud.currentTime = 0;
                                btn.innerHTML = '▶ {lbl_play}'; 
                                btn.style.color = '#3a6a8a';
                                btn.style.borderColor = '#1a2535';
                            }}
                            aud.onended = () => {{ 
                                btn.innerHTML = '▶ {lbl_play}'; 
                                btn.style.color = '#3a6a8a';
                                btn.style.borderColor = '#1a2535';
                            }};
                        }} else {{
                            // Se NÃO tem áudio, usa a voz do navegador do usuário
                            if (window.speechSynthesis.speaking) {{
                                window.speechSynthesis.cancel();
                                btn.innerHTML = '▶ {lbl_play}';
                                btn.style.color = '#3a6a8a';
                                btn.style.borderColor = '#1a2535';
                            }} else {{
                                var textToSpeak = '{clean_text}';
                                var u = new SpeechSynthesisUtterance(textToSpeak);
                                u.lang = 'en-US'; // Lê o texto com sotaque americano
                                u.onend = function() {{
                                    btn.innerHTML = '▶ {lbl_play}';
                                    btn.style.color = '#3a6a8a';
                                    btn.style.borderColor = '#1a2535';
                                }};
                                window.speechSynthesis.speak(u);
                                btn.innerHTML = '⏹ {lbl_stop}';
                                btn.style.color = '#e05c2a';
                                btn.style.borderColor = 'rgba(224,92,42,.5)';
                            }}
                        }}
                    " style="
                        background: transparent; border: 1px solid #1a2535; 
                        color: #3a6a8a; border-radius: 8px; 
                        padding: 4px 12px; cursor: pointer; font-size: 0.75rem;
                        transition: all 0.15s; font-family: sans-serif;
                    ">
                        ▶ {lbl_play}
                    </button>
                </div>
                """
                st.markdown(audio_html, unsafe_allow_html=True)

    # ── Indicador "digitando" — aparece enquanto Claude processa ────
    if st.session_state.get("speaking"):
        components.html("""<!DOCTYPE html><html><head>
<style>
*{margin:0;padding:0;box-sizing:border-box;}
html,body{background:transparent;overflow:hidden;font-family:'Sora',sans-serif;}
.typing-row{
  display:flex;align-items:center;gap:10px;
  padding:6px 0 4px 0;
}
.av{
  width:30px;height:30px;border-radius:50%;
  background:#1e2a3a;display:flex;align-items:center;
  justify-content:center;font-size:14px;flex-shrink:0;
}
.typing-bubble{
  display:flex;align-items:center;gap:8px;
  background:#1a1f2e;border:1px solid #252d3d;
  border-radius:18px;border-bottom-left-radius:4px;
  padding:8px 14px;
}
.spin{
  color:#e05c2a;font-size:16px;line-height:1;
  animation:spinme 1.2s linear infinite;
  display:inline-block;flex-shrink:0;
}
@keyframes spinme{from{transform:rotate(0deg);}to{transform:rotate(360deg);}}
.typing-text{font-size:.75rem;color:#8b949e;font-style:italic;letter-spacing:.2px;}
</style></head><body>
<div class="typing-row">
  <div class="av">🧑‍🏫</div>
  <div class="typing-bubble">
    <span class="spin">✳</span>
    <span class="typing-text" id="msg">Pensando…</span>
  </div>
</div>
<script>
(function(){
  var msgs = [
    [0,   "Pensando…"],
    [3,   "Elaborando resposta…"],
    [7,   "Demorando mais que o normal. Tentando novamente em breve (tentativa 1)"],
    [14,  "Demorando mais que o normal. Tentando novamente em breve (tentativa 2)"],
    [22,  "Demorando mais que o normal. Tentando novamente em breve (tentativa 3)"],
  ];
  var el    = document.getElementById('msg');
  var start = Date.now();
  function update(){
    var sec = (Date.now()-start)/1000;
    var txt = msgs[0][1];
    for(var i=0;i<msgs.length;i++){ if(sec>=msgs[i][0]) txt=msgs[i][1]; }
    el.textContent = txt;
    setTimeout(update, 800);
  }
  update();
})();
</script>
</body></html>""", height=52, scrolling=False)
    staged = st.session_state.get("staged_file")
    if staged:
        # suporta lista (múltiplos) ou dict legado (único)
        staged_list = staged if isinstance(staged, list) else [staged]
        icons = {"audio": "🎵", "text": "📄", "image": "📸"}
        items_html = "".join(
            f'<span style="background:rgba(255,255,255,.06);border-radius:6px;'
            f'padding:3px 8px;font-size:.8rem;color:#e6edf3;">'
            f'{icons.get(f["kind"],"📎")} {f["name"]}</span>'
            for f in staged_list
        )
        st.markdown(f"""
<div style="background:rgba(240,165,0,.08);border:1px solid rgba(240,165,0,.25);
     border-radius:10px;padding:10px 14px;margin:6px 0;
     display:flex;align-items:center;justify-content:space-between;gap:10px;flex-wrap:wrap;">
  <div style="display:flex;gap:6px;flex-wrap:wrap;align-items:center;">
    {items_html}
    <span style="color:#8b949e;font-size:.75rem;">· {len(staged_list)} arquivo(s) anexado(s)</span>
  </div>
  <span style="font-size:.7rem;color:#f0a500;">↩ Digite uma mensagem ou envie</span>
</div>
""", unsafe_allow_html=True)
        if st.button(t("remove_attachment", ui_lang), key="remove_staged"):
            st.session_state.staged_file      = None
            st.session_state.staged_file_name = None
            st.session_state.pop("_last_files_key", None)
            st.rerun()

    # ── Botão de download de arquivo gerado pela IA ───────────────────────────
    pending_dl = st.session_state.get("_pending_download")
    if pending_dl:
        b64_data = pending_dl["b64"]
        fname    = pending_dl["filename"]
        mime     = pending_dl["mime"]
        st.markdown(f"""
<div style="background:rgba(240,165,0,.08);border:1px solid rgba(240,165,0,.35);
     border-radius:10px;padding:10px 16px;margin:8px 0;display:flex;
     align-items:center;justify-content:space-between;gap:12px;">
  <span style="font-size:.85rem;color:#e6edf3;">📎 <b>{fname}</b> pronto para download</span>
  <a href="data:{mime};base64,{b64_data}" download="{fname}"
     style="background:linear-gradient(135deg,#f0a500,#e05c2a);color:#060a10;
     font-weight:700;font-size:.78rem;padding:6px 16px;border-radius:20px;
     text-decoration:none;white-space:nowrap;">⬇ Baixar arquivo</a>
</div>""", unsafe_allow_html=True)

    # ── Chat input (texto) ────────────────────────────────────────────────────
    prompt = st.chat_input(t("type_message", ui_lang))
    if prompt:
        if not API_KEY:
            st.error("Configure ANTHROPIC_API_KEY no .env"); st.stop()

        staged = st.session_state.get("staged_file")
        if staged:
            staged_list = staged if isinstance(staged, list) else [staged]
            for i, sf in enumerate(staged_list):
                extra = prompt if i == 0 else ""
                _process_and_send_file(username, user, conv_id,
                                       sf["raw"], sf["name"], extra_text=extra)
            st.session_state.staged_file      = None
            st.session_state.staged_file_name = None
            st.session_state.pop("_last_files_key", None)
        else:
            append_message(username, conv_id, "user", prompt)
            st.session_state.speaking = True
            try:   send_to_claude(username, user, conv_id, prompt)
            except Exception as e: st.error(f"❌ {e}")
            st.session_state.speaking = False

        st.rerun()

    # ── Gravador de áudio nativo (st.audio_input) ─────────────────────────────
    audio_val = st.audio_input(" ", key=f"voice_input_{st.session_state.audio_key}",
                               label_visibility="collapsed")
    if audio_val and audio_val != st.session_state.get("_last_audio"):
        st.session_state["_last_audio"] = audio_val
        with st.spinner("Transcrevendo..."):
            txt = transcribe_bytes(audio_val.read(), ".wav", None)
        if txt and not txt.startswith("❌") and not txt.startswith("⚠️"):
            if not API_KEY: st.error("Configure ANTHROPIC_API_KEY"); st.stop()
            append_message(username, conv_id, "user", txt, audio=True)
            st.session_state.speaking = True
            try:   send_to_claude(username, user, conv_id, txt)
            except Exception as e: st.error(f"❌ {e}")
            st.session_state.speaking = False
            st.session_state.audio_key += 1
            st.rerun()
        elif txt:
            st.error(txt)

    # ── File uploader (oculto — acionado pelo clipe na chat bar) ─────────────
    uploaded_list = st.file_uploader(
        "📎", key="file_upload", label_visibility="collapsed",
        accept_multiple_files=True,
        type=["mp3", "wav", "ogg", "m4a", "webm", "flac",
              "pdf", "doc", "docx", "txt", "png", "jpg", "jpeg", "webp"])

    if uploaded_list:
        names_key = ",".join(sorted(f.name for f in uploaded_list))
        if names_key != st.session_state.get("_last_files_key"):
            st.session_state["_last_files_key"] = names_key
            staged_list = []
            for uf in uploaded_list:
                raw    = uf.read()
                result = extract_file(raw, uf.name)
                staged_list.append({"raw": raw, "name": uf.name,
                                     "kind": result["kind"], "result": result})
            st.session_state.staged_file      = staged_list
            st.session_state.staged_file_name = ", ".join(f["name"] for f in staged_list)
            st.rerun()

    # ── Handler único para todos os botões Ouvir (substitui iframes por mensagem) ─
    components.html("""<!DOCTYPE html><html><body><script>
(function(){
  var par = window.parent ? window.parent.document : document;
  var cur = null;
  function initBtns(){
    par.querySelectorAll('[data-pav-tts]').forEach(function(btn){
      if(btn._pavInit) return;
      btn._pavInit = true;
      btn.addEventListener('click', function(){
        if(cur && cur !== btn){
          speechSynthesis.cancel();
          cur.textContent='▶ Ouvir'; cur.classList.remove('speaking'); cur=null;
        }
        if(btn.classList.contains('speaking')){
          speechSynthesis.cancel();
          btn.textContent='▶ Ouvir'; btn.classList.remove('speaking'); cur=null; return;
        }
        var txt = btn.getAttribute('data-text') || '';
        var u = new SpeechSynthesisUtterance(txt);
        u.lang='en-US'; u.rate=0.95; u.pitch=1.05;
        speechSynthesis.getVoices();
        setTimeout(function(){
          var vv=speechSynthesis.getVoices();
          var pick=vv.find(function(v){return v.lang==='en-US';})||vv.find(function(v){return v.lang.startsWith('en');});
          if(pick) u.voice=pick;
          u.onstart=function(){ btn.textContent='⏹ Parar'; btn.classList.add('speaking'); cur=btn; };
          u.onend=u.onerror=function(){ btn.textContent='▶ Ouvir'; btn.classList.remove('speaking'); cur=null; };
          speechSynthesis.cancel(); speechSynthesis.speak(u);
        },80);
      });
    });
  }
  initBtns();
  var obs = new MutationObserver(initBtns);
  obs.observe(par.body, {childList:true, subtree:true});
})();
</script></body></html>""", height=1)

    # ── Botões mic e clipe — carregados de arquivos externos ──────────────────
    _btn_html = Path("static/pav_buttons.html")
    _btn_css  = Path("static/pav_buttons.css")
    # Injeta CSS via st.markdown (funciona direto no document)
    if _btn_css.exists():
        st.markdown(f"<style>{_btn_css.read_text()}</style>", unsafe_allow_html=True)
    # Injeta JS via components.html (necessário para executar scripts)
    if _btn_html.exists():
        components.html(_btn_html.read_text(), height=1)
    else:
        st.warning("static/pav_buttons.html não encontrado")


# ══════════════════════════════════════════════════════════════════════════════
# DASHBOARD DA PROFESSORA
# ══════════════════════════════════════════════════════════════════════════════

def show_dashboard() -> None:
    """Painel administrativo com estatísticas de todos os alunos."""
    user    = st.session_state.user
    profile = user.get("profile", {})
    ui_lang = profile.get("language", "pt-BR")

    with st.sidebar:
        st.markdown(f"""<div style="display:flex;align-items:center;gap:10px;margin-bottom:12px;">
            {avatar_html(44)}<div>
            <div style="font-weight:600;font-size:.9rem;">{PROF_NAME}</div>
            <div style="font-size:.7rem;color:#8b949e;"><span class="status-dot"></span>Professora</div>
            </div></div>
            <hr style="border-color:#30363d;margin:6px 0 12px">""", unsafe_allow_html=True)
        if st.button("📊 Dashboard", use_container_width=True, type="primary"): pass
        if st.button(t("voice_mode", ui_lang), use_container_width=True, key="dash_voice"):
            st.session_state.page = "chat"
            st.session_state.voice_mode = True; st.rerun()
        if st.button(t("use_as_student", ui_lang), use_container_width=True, key="dash_chat"):
            st.session_state.page = "chat"; st.rerun()
        if st.button(t("my_profile", ui_lang), use_container_width=True, key="dash_profile"):
            st.session_state.page = "profile"; st.rerun()
        if st.button(t("logout", ui_lang), use_container_width=True, key="dash_logout"):
            _logout(); st.rerun()

    st.markdown("## 📊 Painel do Professor")
    st.markdown("---")
    col_h1, col_h2 = st.columns([4, 1])
    with col_h2:
        if st.button(t("enter_chat", ui_lang), use_container_width=True):
            st.session_state.page = "chat"; st.rerun()
    st.markdown("---")

    stats = get_all_students_stats()
    today = datetime.now().strftime("%Y-%m-%d")
    c1, c2, c3, c4 = st.columns(4)
    for col, val, lbl in zip(
        [c1, c2, c3, c4],
        [len(stats),
         sum(s["messages"]    for s in stats),
         sum(s["corrections"] for s in stats),
         sum(1 for s in stats if s["last_active"][:10] == today)],
        ["Alunos", "Mensagens", "Correções", "Ativos Hoje"]
    ):
        col.markdown(
            f'<div class="stat-card"><div class="val">{val}</div><div class="lbl">{lbl}</div></div>',
            unsafe_allow_html=True)

    st.markdown("<br>")
    st.markdown("### 👥 Alunos")
    if not stats:
        st.info("Nenhum aluno ainda.")
    else:
        badge = {"Beginner": "badge-blue", "Pre-Intermediate": "badge-green",
                 "Intermediate": "badge-gold",   "Business English": "badge-gold"}
        rows = "".join(f"""<tr>
            <td><b>{s['name']}</b><br><span style="color:#8b949e;font-size:.75rem">@{s['username']}</span></td>
            <td><span class="badge {badge.get(s['level'],'badge-blue')}">{s['level']}</span></td>
            <td>{s['focus']}</td>
            <td style="font-family:'JetBrains Mono',monospace;color:#f0a500">{s['messages']}</td>
            <td style="font-family:'JetBrains Mono',monospace;color:#f0a500">{s['corrections']}</td>
            <td style="color:#8b949e">{s['last_active']}</td>
        </tr>""" for s in sorted(stats, key=lambda x: x["messages"], reverse=True))
        st.markdown(
            f'<div style="background:var(--surface);border:1px solid var(--border);'
            f'border-radius:12px;overflow:hidden"><table class="dash-table"><thead>'
            f'<tr><th>Aluno</th><th>Nível</th><th>Foco</th><th>Msgs</th>'
            f'<th>Correções</th><th>Último Acesso</th></tr></thead>'
            f'<tbody>{rows}</tbody></table></div>',
            unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# ROTEADOR PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

if not st.session_state.logged_in:
    show_login()
else:
    # Salva token apenas uma vez por sessão (não a cada rerun)
    _tok = st.session_state.get("_session_token", "")
    if _tok and not st.session_state.get("_session_saved"):
        js_save_session(_tok)
        st.session_state["_session_saved"] = True

    if st.session_state.page == "profile":
        show_profile()
    elif st.session_state.page == "dashboard":
        show_dashboard()
    else:
        show_chat()